"""
Work Summary – Kanini Resume Builder POC
Generates a concise Word document describing the work done.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── colour palette ──────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x6B)
ORANGE = RGBColor(0xE8, 0x6A, 0x1A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF2, 0xF2, 0xF2)
BLACK  = RGBColor(0x00, 0x00, 0x00)


def shade_cell(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = NAVY
    # underline for level-1
    if level == 1:
        run.underline = True
    return p


def add_body(doc, text, size=11, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = BLACK
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    return p


def add_table_2col(doc, rows_data, header=None):
    """Two-column table with optional header row."""
    col_count = 2
    table = doc.add_table(rows=0, cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        hrow = table.add_row()
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            shade_cell(cell, '1B3A6B')
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = WHITE

    for left, right in rows_data:
        row = table.add_row()
        shade_cell(row.cells[0], 'F2F2F2')
        run_l = row.cells[0].paragraphs[0].add_run(left)
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_r = row.cells[1].paragraphs[0].add_run(right)
        run_r.font.size = Pt(11)

    # equal column widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.2)

    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Title banner ──────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after  = Pt(6)
    run = title_para.add_run('KANINI Resume Builder – Work Summary')
    run.bold = True
    run.font.size  = Pt(18)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run(f'Prepared by: Indira Eswaran  |  Date: {datetime.date.today().strftime("%d %B %Y")}')
    r.font.size  = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r.italic = True

    # ── 1. Project Overview ───────────────────────────────────────────
    add_heading(doc, '1. Project Overview')
    add_body(doc,
        'I developed a full-stack Resume Builder POC (Proof of Concept) for Kanini. '
        'The application allows HR teams to upload a candidate\'s resume (PDF or DOCX), '
        'automatically parse the content, and generate professionally formatted Word '
        'documents in two Kanini-approved template styles — ready for client submission.')

    # ── 2. What I Built ───────────────────────────────────────────────
    add_heading(doc, '2. What I Built')

    add_heading(doc, 'A.  AI-Powered Resume Parsing', level=2)
    for b in [
        'Integrated OpenAI GPT-4o-mini to intelligently extract structured data from any resume.',
        'The AI identifies: Candidate Name, Contact Details, Profile Summary, Technical Skills, Work Experience, Projects, and Educational Qualifications.',
        'Implemented a regex-based fallback parser so the app works even without an API key.',
    ]:
        add_bullet(doc, b)

    add_heading(doc, 'B.  Two Professional Resume Templates', level=2)
    add_table_2col(doc, [
        ('Template 1 – Kanini Format ',
         'Calibri font, bold section headers, bullet-list responsibilities, borderless experience table. '
         'Matches the official "Sample KANINI Profile.docx" exactly.'),
        ('Template 2 – Kanini Profile Format',
         'Arial font, right-aligned candidate name, blue underlined "Heading 1" section titles, '
         'tab-aligned fields. Matches "Kanini Format .docx" exactly.'),
    ])

    add_heading(doc, 'C.  Python FastAPI Backend', level=2)
    for b in [
        'Built REST API endpoints: health check, resume upload & parse, and file download.',
        'Session-based file management — each upload gets a unique session ID.',
        'Supports download of generated Word (.docx) and PDF (.pdf) files.',
        'Secure file handling: 10 MB size limit, PDF/DOCX only, no path traversal.',
    ]:
        add_bullet(doc, b)

    add_heading(doc, 'D.  Angular 19 Frontend', level=2)
    for b in [
        'Drag-and-drop file upload with format validation and loading animation.',
        'Side-by-side template preview (live HTML render) scaled to fit the screen.',
        'Candidate name displayed inside each template card.',
        'Download buttons for both Word and PDF formats per template.',
        'Responsive layout (collapses to single column on smaller screens).',
    ]:
        add_bullet(doc, b)

    # ── 3. Technology Stack ───────────────────────────────────────────
    add_heading(doc, '3. Technology Stack')
    add_table_2col(doc, [
        ('Frontend',  'Angular 19, TypeScript, SCSS'),
        ('Backend',   'Python 3.14, FastAPI, Uvicorn'),
        ('AI Parser', 'OpenAI GPT-4o-mini (JSON mode)'),
        ('Documents', 'python-docx 1.1.0, docx2pdf'),
        ('PDF/DOCX Parsing', 'pdfminer.six, docx2txt'),
        ('Config',    '.env file for API key management'),
    ])

    # ── 4. Key Challenges Solved ──────────────────────────────────────
    add_heading(doc, '4. Key Challenges Solved')
    add_table_2col(doc, [
        ('Exact template matching',
         'Reverse-engineered both company DOCX samples paragraph-by-paragraph to reproduce '
         'identical Word structure including soft line breaks, heading styles, and table layouts.'),
        ('Dynamic HTML preview',
         'Generated inline HTML with matching CSS classes so the browser preview '
         'visually mirrors the downloaded Word document.'),
        ('Scalable preview',
         'Used ResizeObserver + CSS transform to scale 780 px wide templates '
         'to fit any screen width without distortion.'),
        ('AI + fallback chain',
         'AI parser runs first; if unavailable or parsing fails, regex parser takes over '
         'transparently — zero downtime for users.'),
        ('CSS architecture',
         'Separated global template-preview styles (styles.scss) from component-scoped '
         'layout styles (component SCSS files) to follow Angular best practices.'),
    ])

    # ── 5. Output ─────────────────────────────────────────────────────
    add_heading(doc, '5. Deliverables')
    for b in [
        'Fully functional web application (Angular + FastAPI).',
        'Two Word document templates matching Kanini\'s exact format requirements.',
        'AI parsing integration with OpenAI GPT-4o-mini.',
        'PDF export capability for both templates.',
        'Clean, maintainable codebase with proper component structure.',
        'This work-summary document.',
    ]:
        add_bullet(doc, b)

    # ── footer note ───────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run('— End of Work Summary —')
    nr.italic = True
    nr.font.size  = Pt(10)
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    out = r'd:\Kanini_Template_POC\Kanini_Work_Summary.docx'
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    build()
