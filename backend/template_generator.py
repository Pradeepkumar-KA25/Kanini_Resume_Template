"""
Kanini Resume Template Generator
Generates two professional Word documents and HTML previews.

Template 1 – Kanini Format    : matches 'Sample KANINI Profile.docx'
Template 2 – Kanini Profile   : matches 'Kanini Format .docx' (Deloitte submission)
"""

import os
import re
import sys
import tempfile
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Brand Colours ────────────────────────────────────────────────────────────
KANINI_DARK    = "1B3A6B"
KANINI_MID     = "2A5298"
KANINI_LIGHT   = "E8F0FB"
KANINI_ACCENT  = "E87722"
KANINI_TEXT    = "2D3748"
KANINI_GRAY    = "718096"
KANINI_SIDEBAR = "2C3E50"
WHITE          = "FFFFFF"

# Heading colour used in FORMAT PROFILE headers
KF_HDR_COLOR   = "0072B4"
BODY_TEXT_COLOR = "6B7280"


def _workspace_root() -> Path:
    return Path(__file__).resolve().parent.parent


def _resolve_logo_path(kind: str) -> Path | None:
    root = _workspace_root()
    kind_low = str(kind or "").strip().lower()
    if kind_low == "kanini":
        candidates = [
            root / "Kanini_Logo_light.png",
            root / "Kanini logo 2.png",
            root / "frontend-ng" / "public" / "kanini-template-logo.png",
            root / "frontend-ng" / "public" / "kanini-logo.png",
            root / "kanini-logo.png",
            root / "Kanini_Logo_Image.png",
            root / "frontend-ng" / "public" / "Kanini_Logo_Image.png",
            root / "download.webp",
            root / "Kanini Logo.webp",
            root / "frontend-ng" / "public" / "kanini-template-icon.webp",
        ]
    elif kind_low == "deloitte":
        candidates = [
            root / "Deloitte_Logo.webp",
            root / "frontend-ng" / "public" / "deloitte-template-icon.webp",
            root / "OIP.webp",
            root / "deloitte.webp",
        ]
    else:
        candidates = []

    existing = [c for c in candidates if c.exists()]
    if not existing:
        return None
    if kind_low == "kanini":
        # The supplied compact wordmark is intentionally the first candidate.
        return existing[0]
    try:
        from PIL import Image

        def _width(path: Path) -> int:
            with Image.open(path) as img:
                return img.width

        return max(existing, key=_width)
    except Exception:
        return existing[0]


def _get_kanini_logo_base64(target_height_px: int = 88) -> str:
    """Return the Kanini logo as a flattened RGB base64 PNG for HTML/PDF.

    Use the small wordmark file in the top-left margin for downloads.
    """
    logo_src = _resolve_logo_path("kanini")
    if not logo_src:
        return ""
    try:
        from PIL import Image
        import base64
        import io

        with Image.open(logo_src) as img:
            if img.mode in ("RGBA", "LA", "P"):
                alpha = img.convert("RGBA").getchannel("A")
                visible_bounds = alpha.getbbox()
                if visible_bounds:
                    img = img.crop(visible_bounds)
            if img.mode in ("RGBA", "LA", "P"):
                background = Image.new("RGB", img.size, (255, 255, 255))
                mask = img.split()[-1] if img.mode != "P" else None
                background.paste(img, mask=mask)
                rgb_img = background
            else:
                rgb_img = img.convert("RGB")

        # Supply more pixels than the fixed CSS display width so the PDF stays sharp.
        ratio = target_height_px / rgb_img.height
        new_size = (int(rgb_img.width * ratio), target_height_px)
        rgb_img = rgb_img.resize(new_size, Image.Resampling.LANCZOS)

        buf = io.BytesIO()
        rgb_img.save(buf, format="PNG")
        encoded = base64.b64encode(buf.getvalue()).decode("utf-8")
        return f"data:image/png;base64,{encoded}"
    except Exception:
        return ""


def _inject_logo_into_html(html_content: str, css_class: str) -> str:
    """Prepend the Kanini logo to a preview HTML block for downloads only."""
    logo_data = _get_kanini_logo_base64()
    if not logo_data:
        return html_content
    logo_html = f'<div class="{css_class}"><img src="{logo_data}" alt="Kanini Logo"></div>'
    # Insert immediately after the opening resume wrapper div.
    first_div_end = html_content.find(">", html_content.find("<div"))
    if first_div_end == -1:
        return html_content
    return html_content[:first_div_end + 1] + "\n  " + logo_html + html_content[first_div_end + 1:]


def _prepare_docx_logo(logo_path: Path) -> Path | None:
    """Ensure image format is compatible with python-docx/Word for headers.

    Word COM PDF conversion can hang when a header contains a PNG with an
    alpha channel (common for WEBP-to-PNG conversions). Flatten transparency
    to a white background to avoid that.
    """
    if not logo_path:
        return None

    suffix = logo_path.suffix.lower()
    is_native_supported = suffix in {".png", ".jpg", ".jpeg", ".bmp", ".gif"}

    # Convert WEBP/other unsupported formats, and also re-save PNGs that may
    # carry an alpha channel, to a flattened RGB PNG.
    try:
        from PIL import Image
        out_dir = Path(tempfile.gettempdir()) / "kanini_resume_builder" / "logo_cache"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{logo_path.stem}_flattened.png"
        with Image.open(logo_path) as img:
            if logo_path.name == "Kanini_Logo_light.png":
                alpha = img.convert("RGBA").getchannel("A")
                visible_bounds = alpha.getbbox()
                if visible_bounds:
                    img = img.crop(visible_bounds)
            if img.mode in ("RGBA", "LA", "P"):
                background = Image.new("RGB", img.size, (255, 255, 255))
                mask = img.split()[-1] if img.mode != "P" else None
                background.paste(img, mask=mask)
                rgb_img = background
            else:
                rgb_img = img.convert("RGB")
            rgb_img.save(out_path, format="PNG")
        return out_path
    except Exception:
        # Fallback: return original path only if it was already a supported format.
        return str(logo_path) if is_native_supported else None


def _add_logo_header(doc: Document, logo_kind: str, width_cm: float = 2.8) -> None:
    logo_src = _resolve_logo_path(logo_kind)
    if not logo_src:
        return

    logo_for_docx = _prepare_docx_logo(logo_src)
    if not logo_for_docx:
        return

    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.clear()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        run.add_picture(str(logo_for_docx), width=Cm(width_cm))


def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ─── XML Helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tcPr.append(shd)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for name, color in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if color:
            el = OxmlElement(f"w:{name}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color.lstrip("#"))
            tcBorders.append(el)


def _add_para_bottom_border(para, color_hex: str):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bottom)


def _set_col_width(table, col_idx: int, width_twips: int):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(width_twips))
        tcW.set(qn("w:type"), "dxa")


# ─── Shared helpers ───────────────────────────────────────────────────────────

def _add_run(para, text: str, bold=False, italic=False, size_pt=10,
             color=None, font_name="Calibri"):
    if font_name == "Times New Roman":
        bold = False
        size_pt = 9.75
        color = BODY_TEXT_COLOR
        para.paragraph_format.line_spacing = 1.6
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = _rgb(color)
    return run


def _add_heading_run(para, text: str, color=KF_HDR_COLOR):
    para.paragraph_format.line_spacing = 1.0
    return _add_run(para, str(text).upper(), bold=True, size_pt=14, color=color, font_name="Inter")


def _add_name_run(para, text: str):
    para.paragraph_format.line_spacing = 1.0
    return _add_run(para, text, bold=True, size_pt=14, color="000000", font_name="Inter")


def _cell_para(cell, text="", bold=False, italic=False, size_pt=10,
               color=None, align=WD_ALIGN_PARAGRAPH.LEFT, font_name="Calibri"):
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    if text:
        _add_run(para, text, bold=bold, italic=italic, size_pt=size_pt,
                 color=color, font_name=font_name)
    return para


def _safe(value, default=""):
    return value if value else default


def _to_roman(n: int) -> str:
    """Convert integer 1-10 to Roman numeral string."""
    vals = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    result = ""
    for v, s in vals:
        while n >= v:
            result += s
            n -= v
    return result


# Leading list markers that may survive extraction: bullet glyphs, dashes,
# arrows, or "1." / "1)" style numbering. Stripped so the template's own
# bullet (the <li> disc or Word list style) is not doubled up.
_LEADING_MARKER_RE = re.compile(
    r"^\s*(?:[•‣◦⁃∙·▪▫●○∙"
    r"–—•▪▫●○◦‣·►▶➤*\-]+\s*|\d+[.)]\s+)"
)


def _clean_bullet_text(text: str) -> str:
    """Remove a single leading bullet/number marker from a line of text."""
    return _LEADING_MARKER_RE.sub("", str(text or "")).strip()


def _split_summary(text: str) -> List[str]:
    """Preserve summary content while normalizing only obvious bullet markers.

    The previous implementation rewrote prose into sentence bullets and added
    trailing punctuation. That changed the user's content. This version keeps
    line breaks if they exist, otherwise returns the original summary as one
    block so the template preserves the source wording more faithfully.
    """
    raw = str(text or "").strip()
    if not raw:
        return []

    # Normalize common bullet glyph variants from PDF/Word extraction.
    normalized = (
        raw.replace("\uf0b7", "•")  # private-use bullet often seen from DOC/PDF extraction
           .replace("", "•")
           .replace("", "•")
    )

    # Normalize line endings and split explicit lines first.
    lines = [ln.strip() for ln in normalized.replace("\r\n", "\n").replace("\r", "\n").split("\n") if ln.strip()]
    if not lines:
        return []

    # Split inline bullet markers as well, for inputs like:
    # "• Point A • Point B • Point C"
    inline_bullet_re = re.compile(r"\s*[•\u2022\u2023\u25E6\u2043\u2219\u2192\u25BA]+\s*")
    numbered_bullet_re = re.compile(r"\s+\d+[.)]\s+")

    points: List[str] = []
    for line in lines:
        has_inline_bullets = bool(inline_bullet_re.search(line))
        has_numbered_bullets = bool(numbered_bullet_re.search(line))

        if has_inline_bullets or has_numbered_bullets:
            parts = [line]
            if has_inline_bullets:
                parts = [p for chunk in parts for p in inline_bullet_re.split(chunk)]
            if has_numbered_bullets:
                parts = [p for chunk in parts for p in numbered_bullet_re.split(chunk)]
            for part in parts:
                cleaned = _clean_bullet_text(part)
                if cleaned:
                    points.append(cleaned)
            continue

        cleaned = _clean_bullet_text(line)
        if cleaned:
            points.append(cleaned)

    return points if points else [raw]


_EXP_RENDER_NOISE_RE = re.compile(
    r"\b(?:company\s+name|designation|duration|project\s+summary|roles?\s+and\s+responsibilities|working\s+experience)\b",
    re.IGNORECASE,
)


def _clean_company_for_render(company: str) -> str:
    text = str(company or "").strip().strip(".")
    if not text:
        return ""
    text = re.split(r",\s+a\s+company\s+that\b|,\s+which\s+|\s+speciali[sz]es\s+in\b", text, maxsplit=1, flags=re.IGNORECASE)[0].strip()
    return text.rstrip(",")


def _is_meaningful_exp_value(value: str) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    if _EXP_RENDER_NOISE_RE.search(text):
        return False
    if text in {":", "-"}:
        return False
    return True


def _prepare_experience_for_render(exp_list: Any) -> List[Dict[str, Any]]:
    if not isinstance(exp_list, list):
        return []

    prepared: List[Dict[str, Any]] = []
    for raw in exp_list:
        if not isinstance(raw, dict):
            continue

        company = _clean_company_for_render(str(raw.get("company") or "").strip())
        company_name = _clean_company_for_render(str(raw.get("company_name") or raw.get("company") or "").strip())
        title = str(raw.get("title") or "").strip().strip(".")
        dates = str(raw.get("dates") or "").strip().strip(".")
        responsibilities = raw.get("responsibilities") if isinstance(raw.get("responsibilities"), list) else []

        meaningful_count = sum(
            1 for value in (company, title, dates) if _is_meaningful_exp_value(value)
        )
        if meaningful_count < 2:
            continue

        current = {
            "company": company,
            "company_name": company_name or company,
            "company_sector": str(raw.get("company_sector") or "").strip(),
            "title": title,
            "dates": dates,
            "location": str(raw.get("location") or "").strip(),
            "responsibilities": responsibilities,
        }

        merged = False
        for existing in prepared:
            same_company = company and existing.get("company") and company.casefold() == str(existing.get("company") or "").casefold()
            same_title = title and existing.get("title") and title.casefold() == str(existing.get("title") or "").casefold()
            if same_company or same_title:
                if not existing.get("company") and company:
                    existing["company"] = company
                if not existing.get("title") and title:
                    existing["title"] = title
                if not existing.get("dates") and dates:
                    existing["dates"] = dates
                if not existing.get("company_name") and current.get("company_name"):
                    existing["company_name"] = current["company_name"]
                if not existing.get("company_sector") and current.get("company_sector"):
                    existing["company_sector"] = current["company_sector"]
                if responsibilities:
                    seen = {str(r).casefold() for r in existing.get("responsibilities", [])}
                    for resp in responsibilities:
                        if str(resp).casefold() not in seen:
                            existing.setdefault("responsibilities", []).append(resp)
                merged = True
                break

        if not merged:
            prepared.append(current)

    return prepared


def _template1_company_value(exp: Dict[str, Any]) -> str:
    """Format1 should always show company names from resume, not sector labels."""
    name = _clean_company_for_render(str((exp or {}).get("company_name") or "").strip())
    if name:
        return name
    return _clean_company_for_render(str((exp or {}).get("company") or "").strip())


def _template2_sector_value(exp: Dict[str, Any]) -> str:
    company = str((exp or {}).get("company") or "").strip()
    sector = str((exp or {}).get("company_sector") or "").strip()

    # Keep sector-only rendering for template2; do not leak employer names.
    if sector and sector.casefold() != company.casefold():
        return sector
    return "Industry Not Determined"


def _add_br(para):
    """Insert a soft line break (w:br) within a paragraph."""
    run = para.add_run()
    br = OxmlElement('w:br')
    run._r.append(br)


# ═══════════════════════════════════════════════════════════════════════════════
# TEMPLATE 1 — Kanini Format   (matches Sample KANINI Profile.docx)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_template1(data: Dict, output_path: str):
    """
    Kanini Format  resume – structure mirrors 'Sample KANINI Profile.docx':
      • CANDIDATE NAME bold at top
      • Contact info
      • Profile Summary (bold hdr) + bullet list
      • Technical Skills: (bold hdr) + bullets "Category: items"
      • Work Experience: (bold hdr) + 3-row × 2-col table per job
      • Project Summary: (bold hdr) + Project I/II/III with role|company + resp bullets
      • Educational Qualification: (bold hdr) + text
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21.6)
    sec.page_height = Cm(27.9)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)
    _add_logo_header(doc, "kanini", width_cm=2.0)

    contact  = data.get("contact", {})
    name     = _safe(contact.get("name"), "CANDIDATE NAME")
    exp_list = _prepare_experience_for_render(data.get("experience", []))
    proj_list = data.get("projects", [])
    skills   = data.get("skills", {})
    summary_text = data.get("summary", "")
    education    = data.get("education", [])
    certs        = data.get("certifications", [])
    achievements = data.get("achievements", [])

    F = "Times New Roman"

    # ── Name ────────────────────────────────────────────────────────────────────
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_name_run(np_, name.upper())

    # ── Contact line ─────────────────────────────────────────────────────────────
    contact_parts = []
    if contact.get("phone"):    contact_parts.append(f"Mobile No: {contact['phone']}")
    if contact.get("email"):    contact_parts.append(f"Email Id: {contact['email']}")
    if contact.get("location"): contact_parts.append(contact["location"])
    if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
    if contact.get("github"):   contact_parts.append(contact["github"])
    if contact_parts:
        cp = doc.add_paragraph()
        _add_run(cp, "   |   ".join(contact_parts), size_pt=11, font_name=F)

    doc.add_paragraph()  # spacer

    # ── Professional Summary (fixed section order) ──────────────────────────────
    sh = doc.add_paragraph()
    _add_heading_run(sh, "PROFESSIONAL SUMMARY")
    doc.add_paragraph()
    if summary_text:
        for pt in _split_summary(summary_text):
            bp = doc.add_paragraph(style="List Paragraph")
            _add_run(bp, pt, size_pt=11, font_name=F)
    else:
        ep = doc.add_paragraph()
        _add_run(ep, "-", size_pt=11, font_name=F)
    doc.add_paragraph()

    # ── Technical Skills (fixed section order) ──────────────────────────────────
    sh = doc.add_paragraph()
    _add_heading_run(sh, "Technical Skills:")
    doc.add_paragraph()
    if skills:
        for cat, items in skills.items():
            if items:
                bp = doc.add_paragraph(style="List Paragraph")
                # Entire line bold to match sample
                _add_run(bp, f"{cat}: {', '.join(items)}", bold=True, size_pt=11, font_name=F)
    else:
        ep = doc.add_paragraph()
        _add_run(ep, "-", size_pt=11, font_name=F)
    doc.add_paragraph()

    # ── Work Experience (fixed section order) ───────────────────────────────────
    sh = doc.add_paragraph()
    _add_heading_run(sh, "Work Experience:")
    doc.add_paragraph()

    if exp_list:
        for exp in exp_list:
            # Use fixed-width rows instead of tabs to keep alignment stable in Word/PDF.
            kv_table = doc.add_table(rows=0, cols=3)
            kv_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            kv_table.autofit = False
            _remove_table_borders(kv_table)
            _set_col_width(kv_table, 0, 1700)
            _set_col_width(kv_table, 1, 220)
            _set_col_width(kv_table, 2, 7300)

            for label, value in [
                ("Company Name", _safe(_template1_company_value(exp))),
                ("Designation", _safe(exp.get("title"))),
                ("Duration", _safe(exp.get("dates"))),
            ]:
                row = kv_table.add_row().cells
                _cell_para(row[0], label, size_pt=11, font_name=F)
                _cell_para(row[1], ":", size_pt=11, font_name=F)
                _cell_para(row[2], value or "-", size_pt=11, font_name=F)
                for idx in (0, 1, 2):
                    row[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    row[idx].paragraphs[0].paragraph_format.space_before = Pt(0)
                    row[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
            
            # Responsibilities (bullet points)
            responsibilities = exp.get("responsibilities", [])
            if responsibilities:
                doc.add_paragraph()
                for resp in responsibilities:
                    resp_text = _clean_bullet_text(resp)
                    if resp_text:
                        bp = doc.add_paragraph(style="List Paragraph")
                        bp.paragraph_format.left_indent = Inches(0.35)
                        _add_run(bp, resp_text, size_pt=11, font_name=F)
            
            doc.add_paragraph()  # gap between companies
    else:
        ep = doc.add_paragraph()
        _add_run(ep, "-", size_pt=11, font_name=F)
        doc.add_paragraph()

    # ── Project Summary (fixed section order) ───────────────────────────────────
    # Keep this section after Work Experience in all cases.
    psh = doc.add_paragraph()
    _add_heading_run(psh, "Project Summary:")
    doc.add_paragraph()

    project_entries = proj_list if proj_list else exp_list
    if project_entries:
        for i, project in enumerate(project_entries):
            title_str = _safe(project.get("name") or project.get("title"))
            company_str = _safe(project.get("client") or _template1_company_value(project))
            roman = _to_roman(i + 1)

            # Project label
            pn = doc.add_paragraph()
            _add_heading_run(pn, f"Project {roman}:")

            # Role | Company line
            rc = doc.add_paragraph()
            _add_heading_run(rc, title_str or f"Project {roman}")

            if company_str:
                cp = doc.add_paragraph()
                _add_run(cp, f"Client: {company_str}", size_pt=11, font_name=F)
            if project.get("technologies"):
                tp = doc.add_paragraph()
                _add_run(tp, f"Technologies: {', '.join(project['technologies'])}", size_pt=11, font_name=F)
            if project.get("description"):
                dp = doc.add_paragraph()
                _add_run(dp, project["description"], size_pt=11, font_name=F)

            # Responsibilities heading
            rh = doc.add_paragraph()
            _add_heading_run(rh, "Roles and Responsibilities:")
            doc.add_paragraph()

            for resp in project.get("responsibilities", []):
                resp = _clean_bullet_text(resp)
                if not resp:
                    continue
                bp = doc.add_paragraph(style="List Paragraph")
                bp.paragraph_format.left_indent = Inches(0.35)
                _add_run(bp, resp, size_pt=11, font_name=F)
            doc.add_paragraph()
    else:
        ep = doc.add_paragraph()
        _add_run(ep, "-", size_pt=11, font_name=F)
        doc.add_paragraph()

    # ── Educational Qualification ─────────────────────────────────────────────────
    if education:
        sh = doc.add_paragraph()
        _add_heading_run(sh, "Educational Qualification:")
        doc.add_paragraph()
        for edu in education:
            parts = []
            if edu.get("degree"):      parts.append(edu["degree"])
            if edu.get("year"):        parts.append(f"({edu['year']})")
            if edu.get("institution"): parts.append(f"from {edu['institution']}")
            if edu.get("gpa"):         parts.append(f"GPA: {edu['gpa']}")
            ep = doc.add_paragraph()
            _add_run(ep, " ".join(parts), size_pt=11, font_name=F)

    # ── Certifications ────────────────────────────────────────────────────────────
    if certs:
        doc.add_paragraph()
        sh = doc.add_paragraph()
        _add_heading_run(sh, "Certifications:")
        doc.add_paragraph()
        for cert in certs:
            cert = _clean_bullet_text(cert)
            if not cert:
                continue
            bp = doc.add_paragraph(style="List Paragraph")
            _add_run(bp, cert, size_pt=11, font_name=F)

    if achievements:
        sh = doc.add_paragraph()
        _add_heading_run(sh, "Achievements:")
        for achievement in achievements:
            text = _clean_bullet_text(achievement)
            if text:
                bp = doc.add_paragraph(style="List Paragraph")
                _add_run(bp, text, size_pt=11, font_name=F)

    doc.save(output_path)


# ═══════════════════════════════════════════════════════════════════════════════
# TEMPLATE 2 — Kanini Profile Format  (matches Kanini Format .docx)
#              Used for client submissions (e.g., Deloitte)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_template_deloitte(data: Dict, output_path: str):
    """
    Kanini Profile Format resume – structure mirrors 'Kanini Format .docx':
      • Name right-aligned, bold, 10pt
      • Professional Summary: (Heading 1) + body 12pt
      • Technical Skills: (Heading 1) + tab-separated rows
      • Working Experience: (Heading 1) + tab-separated rows
      • Project Summary: (Heading 1) + Project – I + Client/Technologies/Role + Description + Responsibilities
      • EDUCATIONAL QUALIFICATION: (Heading 1) + text
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.93)
    sec.bottom_margin = Cm(0.49)
    sec.left_margin   = Cm(1.38)
    sec.right_margin  = Cm(0.53)
    _add_logo_header(doc, "kanini", width_cm=2.0)

    contact  = data.get("contact", {})
    name     = _safe(contact.get("name"), "Candidate Name")
    exp_list = _prepare_experience_for_render(data.get("experience", []))
    projects = data.get("projects", [])
    skills   = data.get("skills", {})
    summary_text = data.get("summary", "")
    education    = data.get("education", [])
    certs        = data.get("certifications", [])
    achievements = data.get("achievements", [])

    F = "Times New Roman"

    def add_h1(title: str):
        p = doc.add_paragraph(style="Heading 1")
        _add_heading_run(p, title, color=KF_HDR_COLOR)
        return p

    def add_body(text: str, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        _add_run(p, text, bold=bold, size_pt=12, font_name=F)
        return p

    def add_tab_row(label: str, value: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        _add_run(p, f"{label}\t: {value}", size_pt=12, font_name=F)
        return p

    contact_parts = [
        str(contact[key]).strip()
        for key in ("email", "phone", "location", "linkedin", "github")
        if contact.get(key)
    ]

    # ── Name (center-aligned, 10pt bold) ─────────────────────────────────────────
    doc.add_paragraph()
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_name_run(np_, name.upper())
    if contact_parts:
        add_body(" | ".join(contact_parts))
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Professional Summary ──────────────────────────────────────────────────────
    add_h1("Professional Summary:")
    doc.add_paragraph()
    summary_points = _split_summary(summary_text)
    if summary_points:
        for pt in summary_points:
            add_body(pt)
    else:
        add_body("-")
    doc.add_paragraph()

    # ── Technical Skills ─────────────────────────────────────────────────────────
    add_h1("Technical Skills:")
    doc.add_paragraph()
    rendered_skill_rows = 0
    for cat, items in skills.items():
        if items:
            add_tab_row(cat, ", ".join(items))
            rendered_skill_rows += 1
    if rendered_skill_rows == 0:
        add_tab_row("Skills", "-")
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Working Experience ────────────────────────────────────────────────────────
    add_h1("Working Experience:")
    doc.add_paragraph()
    if exp_list:
        for exp in exp_list:
            add_tab_row("Company Name", _template2_sector_value(exp))
            add_tab_row("Designation",  _safe(exp.get("title"), "-"))
            add_tab_row("Duration",     _safe(exp.get("dates"), "-"))
    else:
        add_tab_row("Company Name", "-")
        add_tab_row("Designation", "-")
        add_tab_row("Duration", "-")
    doc.add_paragraph()

    # ── Project Summary ───────────────────────────────────────────────────────────
    # Build project entries: prefer 'projects' list; fall back to 'experience'
    all_techs_str = ", ".join(item for items in skills.values() for item in items)

    proj_entries: List[Dict] = []
    if projects:
        base_exp = exp_list[0] if exp_list else {}
        for proj in projects:
            techs = proj.get("technologies", [])
            tech_str = techs if isinstance(techs, str) else ", ".join(techs) or all_techs_str
            project_resps = proj.get("responsibilities", []) or []
            base_resps = base_exp.get("responsibilities", []) if isinstance(base_exp, dict) else []
            proj_entries.append({
                "name":            _safe(proj.get("name")),
                "client":          _safe(proj.get("client"), _template2_sector_value(base_exp)),
                "technologies":    tech_str,
                "role":            _safe(proj.get("role"), _safe(base_exp.get("title"))),
                "description":     _safe(proj.get("description")) or (project_resps[0] if project_resps else ""),
                "responsibilities": project_resps or base_resps,
            })
    elif exp_list:
        for exp in exp_list:
            resps = exp.get("responsibilities", [])
            proj_entries.append({
                "name":            "",
                "client":          _template2_sector_value(exp),
                "technologies":    all_techs_str[:300],
                "role":            _safe(exp.get("title")),
                "description":     resps[0] if resps else "",
                "responsibilities": resps[1:] if len(resps) > 1 else resps,
            })

    add_h1("Project Summary:")
    doc.add_paragraph()
    if proj_entries:
        for i, entry in enumerate(proj_entries):
            pn = doc.add_paragraph()
            project_label = entry["name"] or f"Project \u2013 {_to_roman(i + 1)}"
            _add_heading_run(pn, project_label)

            add_tab_row("Client",       entry["client"] or "-")
            add_tab_row("Technologies", entry["technologies"] or "-")
            add_tab_row("Role",         entry["role"] or "-")

            add_h1("Description of Project:")
            add_body(entry["description"] or "-")

            add_h1("Roles and Responsibilities:")
            rendered_resp = False
            for resp in entry["responsibilities"]:
                resp = _clean_bullet_text(resp)
                if resp:
                    add_body(resp)
                    rendered_resp = True
            if not rendered_resp:
                add_body("-")
    else:
        pn = doc.add_paragraph()
        _add_heading_run(pn, "Project \u2013 I")
        add_tab_row("Client", "-")
        add_tab_row("Technologies", "-")
        add_tab_row("Role", "-")
        add_h1("Description of Project:")
        add_body("-")
        add_h1("Roles and Responsibilities:")
        add_body("-")
    doc.add_paragraph()

    # ── Educational Qualification ─────────────────────────────────────────────────
    add_h1("EDUCATIONAL QUALIFICATION:")
    doc.add_paragraph()
    if education:
        for edu in education:
            parts = []
            if edu.get("degree"):      parts.append(edu["degree"])
            if edu.get("year"):        parts.append(f"({edu['year']})")
            if edu.get("institution"): parts.append(f"from {edu['institution']}")
            if edu.get("gpa"):         parts.append(f"GPA: {edu['gpa']}")
            add_body(" ".join(filter(None, parts)) or "-")
    else:
        add_body("-")

    # ── Certifications ────────────────────────────────────────────────────────────
    if certs:
        add_h1("Certifications:")
        doc.add_paragraph()
        for cert in certs:
            cert = _clean_bullet_text(cert)
            if cert:
                add_body(f"\u2022 {cert}")

    if achievements:
        add_h1("Achievements:")
        for achievement in achievements:
            text = _clean_bullet_text(achievement)
            if text:
                add_body(f"\u2022 {text}")

    doc.save(output_path)


# ═══════════════════════════════════════════════════════════════════════════════
# HTML Preview Generators
# ═══════════════════════════════════════════════════════════════════════════════

def _esc(text: str) -> str:
    import html
    return html.escape(str(text))


def generate_preview_html_template1(data: Dict) -> str:
    """HTML preview for Kanini Format  (matches Sample KANINI Profile layout)."""
    contact  = data.get("contact", {})
    name     = _safe(contact.get("name"), "CANDIDATE NAME")
    exp_list = _prepare_experience_for_render(data.get("experience", []))
    proj_list = data.get("projects", [])
    skills   = data.get("skills", {})
    summary_text = data.get("summary", "")
    education    = data.get("education", [])
    certs        = data.get("certifications", [])
    achievements = data.get("achievements", [])

    contact_parts = []
    if contact.get("phone"):
        contact_parts.append(f"Mobile No: {_esc(contact['phone'])}")
    if contact.get("email"):
        contact_parts.append(f"Email Id: {_esc(contact['email'])}")
    if contact.get("location"):
        contact_parts.append(_esc(contact["location"]))
    if contact.get("linkedin"):
        contact_parts.append(_esc(contact["linkedin"]))
    if contact.get("github"):
        contact_parts.append(_esc(contact["github"]))
    contact_str = "   |   ".join(contact_parts)

    # 1) Professional Summary
    if summary_text:
        summary_bullets = "".join(
            f'<li class="t1-bullet">{_esc(pt)}</li>'
            for pt in _split_summary(summary_text)
        )
        summary_content = f'<ul class="t1-bullet-list">{summary_bullets}</ul>'
    else:
        summary_content = '<p class="t1-body">-</p>'
    summary_html = f"""
  <div class="t1-section-hdr">PROFESSIONAL SUMMARY</div>
  <div class="t1-spacer"></div>
  {summary_content}"""

    # 2) Technical Skills
    skills_bullets = "".join(
        f'<li class="t1-bullet"><strong>{_esc(cat)}: {_esc(", ".join(items))}</strong></li>'
        for cat, items in skills.items() if items
    ) if skills else ""
    skills_content = (
        f'<ul class="t1-bullet-list">{skills_bullets}</ul>'
        if skills_bullets else '<p class="t1-body">-</p>'
    )
    skills_html = f"""
  <div class="t1-section-hdr">Technical Skills:</div>
  <div class="t1-spacer"></div>
  {skills_content}"""

    # 3) Work Experience
    exp_entries = ""
    if exp_list:
        for exp in exp_list:
            responsibilities = exp.get("responsibilities", [])
            resp_bullets = "".join(
                f'<li class="t1-bullet">{_esc(_clean_bullet_text(r))}</li>'
                for r in responsibilities if _clean_bullet_text(r)
            )
            exp_entries += f"""
  <table class="t1-exp-table"><tr><td>
  <div class="t1-exp-entry">
        <div class="t1-exp-row"><span class="t1-exp-label">Company Name</span><span class="t1-exp-sep">:</span><span class="t1-exp-val">{_esc(_safe(_template1_company_value(exp)))}</span></div>
    <div class="t1-exp-row"><span class="t1-exp-label">Designation</span><span class="t1-exp-sep">:</span><span class="t1-exp-val">{_esc(_safe(exp.get("title")))}</span></div>
    <div class="t1-exp-row"><span class="t1-exp-label">Duration</span><span class="t1-exp-sep">:</span><span class="t1-exp-val">{_esc(_safe(exp.get("dates")))}</span></div>
    {f'<ul class="t1-bullet-list">{resp_bullets}</ul>' if resp_bullets else ''}
  </div>
  </td></tr></table>"""
    exp_content = exp_entries if exp_entries else '<p class="t1-body">-</p>'
    exp_html = f"""
  <div class="t1-section-hdr">Work Experience:</div>
  <div class="t1-spacer"></div>
  {exp_content}"""

    # 4) Project Summary
    source_projects = proj_list if proj_list else exp_list
    if source_projects and all(not (p.get("description") or p.get("technologies") or p.get("responsibilities")) for p in source_projects):
        source_projects = exp_list
    if source_projects:
        entries = ""
        for i, exp in enumerate(source_projects):
            role_source = exp_list[i] if i < len(exp_list) else {}
            resps = "".join(
                f'<li class="t1-bullet">{_esc(_clean_bullet_text(r))}</li>'
                for r in exp.get("responsibilities", []) if _clean_bullet_text(r)
            )
            roman = _to_roman(i + 1)
            title_str = _safe(exp.get("name") or exp.get("title") or role_source.get("title"))
            company_str = _safe(exp.get("client") or _template1_company_value(exp) or _template1_company_value(role_source))
            if not company_str and exp.get('technologies'):
                company_str = ", ".join(exp.get('technologies', []))
            if not resps and role_source.get("responsibilities"):
                resps = "".join(
                    f'<li class="t1-bullet">{_esc(_clean_bullet_text(r))}</li>'
                    for r in role_source.get("responsibilities", []) if _clean_bullet_text(r)
                )
            entries += f"""
  <div class="t1-spacer"></div>
  <div class="t1-proj-num">Project {roman}:</div>
  <div class="t1-spacer"></div>
  <div class="t1-role-co">{_esc(title_str)} | {_esc(company_str)}</div>
    {f'<p class="t1-body">Technologies: {_esc(", ".join(exp.get("technologies", [])))}</p>' if exp.get("technologies") else ''}
    {f'<p class="t1-body">{_esc(exp.get("description"))}</p>' if exp.get("description") else ''}
  <div class="t1-resp-hdr">Roles and Responsibilities:</div>
  <div class="t1-spacer"></div>
  {('<ul class="t1-bullet-list">' + resps + '</ul>') if resps else '<p class="t1-body">-</p>'}"""
        proj_content = entries
    else:
        proj_content = '<p class="t1-body">-</p>'
    proj_html = f"""
  <div class="t1-section-hdr">Project Summary:</div>
  <div class="t1-spacer"></div>
  {proj_content}"""

    # Additional sections after core Kanini order
    edu_html = ""
    if education:
        items = ""
        for edu in education:
            parts = [_safe(edu.get("degree"))]
            if edu.get("year"):
                parts.append(f"({edu['year']})")
            if edu.get("institution"):
                parts.append(f"from {edu['institution']}")
            if edu.get("gpa"):
                parts.append(f"GPA: {edu['gpa']}")
            items += f'<p class="t1-body">{_esc(" ".join(filter(None, parts)))}</p>'
        edu_html = f"""
  <div class="t1-section-hdr">Educational Qualification:</div>
  <div class="t1-spacer"></div>
  {items}"""

    cert_html = ""
    if certs:
        bullets = "".join(
            f'<li class="t1-bullet">{_esc(_clean_bullet_text(c))}</li>'
            for c in certs if _clean_bullet_text(c)
        )
        cert_html = f"""
  <div class="t1-section-hdr">Certifications:</div>
  <div class="t1-spacer"></div>
  <ul class="t1-bullet-list">{bullets}</ul>"""

    achievement_html = ""
    if achievements:
        bullets = "".join(
            f'<li class="t1-bullet">{_esc(_clean_bullet_text(item))}</li>'
            for item in achievements if _clean_bullet_text(item)
        )
        achievement_html = f"""
  <div class="t1-section-hdr">Achievements:</div>
  <div class="t1-spacer"></div>
  <ul class="t1-bullet-list">{bullets}</ul>"""

    name_header = ""
    if name and name != "CANDIDATE NAME":
        name_header = f'<p class="t1-name">{_esc(name.upper())}</p>'
        if contact_str:
            name_header += f'<p class="t1-contact">{contact_str}</p>'
        name_header += '<div class="t1-spacer"></div>'

    return f"""<div class="t1-resume">
  {name_header}
  {summary_html}
  {skills_html}
  {exp_html}
  {proj_html}
  {edu_html}
  {cert_html}
    {achievement_html}
</div>"""


def generate_preview_html_deloitte(data: Dict) -> str:
    """HTML preview for Kanini Profile Format (Deloitte submission style)."""
    contact  = data.get("contact", {})
    name     = _safe(contact.get("name"), "Candidate Name")
    exp_list = _prepare_experience_for_render(data.get("experience", []))
    projects = data.get("projects", [])
    skills   = data.get("skills", {})
    summary_text = data.get("summary", "")
    education    = data.get("education", [])
    certs        = data.get("certifications", [])
    achievements = data.get("achievements", [])

    def kf_h1(title: str) -> str:
        return f'<div class="kf-h1">{_esc(title.upper())}</div>'

    def kf_tab(label: str, value: str) -> str:
        return f'<p class="kf-body kf-tab-row"><strong>{_esc(label)} :</strong> {_esc(value)}</p>'

    contact_parts = [
        _esc(contact[key])
        for key in ("email", "phone", "location", "linkedin", "github")
        if contact.get(key)
    ]

    summary_points = _split_summary(summary_text)
    if summary_points:
        pts = "".join(
            f'<p class="kf-body">{_esc(pt)}</p>'
            for pt in summary_points
        )
    else:
        pts = '<p class="kf-body">-</p>'
    summary_html = kf_h1("Professional Summary:") + '<div class="kf-spacer"></div>' + pts + '<div class="kf-spacer"></div>'

    skills_bullets = "".join(
        f'<li class="kf-bullet"><strong>{_esc(cat)}: {_esc(", ".join(items))}</strong></li>'
        for cat, items in skills.items() if items
    )
    skills_content = (
        f'<ul class="kf-bullet-list">{skills_bullets}</ul>'
        if skills_bullets
        else '<p class="kf-body">-</p>'
    )
    skills_html = kf_h1("Technical Skills:") + '<div class="kf-spacer"></div>' + skills_content + '<div class="kf-spacer"></div><div class="kf-spacer"></div>'

    if exp_list:
        rows = ""
        for exp in exp_list:
            rows += kf_tab("Company Name", _template2_sector_value(exp))
            rows += kf_tab("Designation",  _safe(exp.get("title"), "-"))
            rows += kf_tab("Duration",     _safe(exp.get("dates"), "-"))
    else:
        rows = kf_tab("Company Name", "-") + kf_tab("Designation", "-") + kf_tab("Duration", "-")
    exp_html = kf_h1("Working Experience:") + '<div class="kf-spacer"></div>' + rows + '<div class="kf-spacer"></div>'

    all_techs_str = ", ".join(item for items in skills.values() for item in items)
    proj_entries: List[Dict] = []
    if projects:
        base_exp = exp_list[0] if exp_list else {}
        for proj in projects:
            techs = proj.get("technologies", [])
            tech_str = techs if isinstance(techs, str) else ", ".join(techs) or all_techs_str
            project_resps = proj.get("responsibilities", []) or []
            base_resps = base_exp.get("responsibilities", []) if isinstance(base_exp, dict) else []
            proj_entries.append({
                "name": _safe(proj.get("name")),
                "client": _safe(proj.get("client"), _template2_sector_value(base_exp)),
                "technologies": tech_str,
                "role": _safe(proj.get("role"), _safe(base_exp.get("title"))),
                "description": _safe(proj.get("description")) or (project_resps[0] if project_resps else ""),
                "responsibilities": project_resps or base_resps,
            })
    elif exp_list:
        for exp in exp_list:
            resps = exp.get("responsibilities", [])
            proj_entries.append({
                "name": "",
                "client": _template2_sector_value(exp),
                "technologies": all_techs_str[:300],
                "role": _safe(exp.get("title")),
                "description": resps[0] if resps else "",
                "responsibilities": resps[1:] if len(resps) > 1 else resps,
            })

    if proj_entries:
        entries_html = ""
        for i, entry in enumerate(proj_entries):
            resps_html = "".join(
                f'<p class="kf-body kf-resp-bullet">&bull;&nbsp;{_esc(_clean_bullet_text(r))}</p>'
                for r in entry["responsibilities"] if _clean_bullet_text(r)
            )
            if not resps_html:
                resps_html = '<p class="kf-body">-</p>'
            entries_html += f"""
    <p class="kf-body kf-proj-num"><strong>{_esc(entry["name"] or f"Project – {_to_roman(i + 1)}")}</strong></p>
  {kf_tab("Client", entry["client"] or "-")}
  {kf_tab("Technologies", entry["technologies"] or "-")}
  {kf_tab("Role", entry["role"] or "-")}
  {kf_h1("Description of Project:")}
  <p class="kf-body">{_esc(entry["description"] or "-")}</p>
  {kf_h1("Roles and Responsibilities:")}
  {resps_html}"""
    else:
        entries_html = f"""
  <p class="kf-body kf-proj-num"><strong>Project &ndash; I</strong></p>
  {kf_tab("Client", "-")}
  {kf_tab("Technologies", "-")}
  {kf_tab("Role", "-")}
  {kf_h1("Description of Project:")}
  <p class="kf-body">-</p>
  {kf_h1("Roles and Responsibilities:")}
  <p class="kf-body">-</p>"""
    proj_html = kf_h1("Project Summary:") + '<div class="kf-spacer"></div>' + entries_html + '<div class="kf-spacer"></div>'

    if education:
        items = ""
        for edu in education:
            parts = [_safe(edu.get("degree"))]
            if edu.get("year"):
                parts.append(f"({edu['year']})")
            if edu.get("institution"):
                parts.append(f"from {edu['institution']}")
            if edu.get("gpa"):
                parts.append(f"GPA: {edu['gpa']}")
            items += f'<p class="kf-body">{_esc(" ".join(filter(None, parts)) or "-")}</p>'
    else:
        items = '<p class="kf-body">-</p>'
    edu_html = kf_h1("EDUCATIONAL QUALIFICATION:") + '<div class="kf-spacer"></div>' + items

    cert_html = ""
    if certs:
        items = "".join(
            f'<p class="kf-body">&bull; {_esc(_clean_bullet_text(c))}</p>'
            for c in certs if _clean_bullet_text(c)
        )
        cert_html = kf_h1("Certifications:") + '<div class="kf-spacer"></div>' + items

    achievement_html = ""
    if achievements:
        items = "".join(
            f'<p class="kf-body">&bull; {_esc(_clean_bullet_text(item))}</p>'
            for item in achievements if _clean_bullet_text(item)
        )
        achievement_html = kf_h1("Achievements:") + '<div class="kf-spacer"></div>' + items

    name_header = ""
    if name and name != "Candidate Name":
        name_header = f'<p class="kf-name">{_esc(name.upper())}</p>'
        if contact_parts:
            name_header += f'<p class="kf-body">{" | ".join(contact_parts)}</p>'

    return f"""<div class="kf-resume">
  <div class="kf-spacer"></div>
  {name_header}
  <div class="kf-spacer"></div>
  {summary_html}
  {skills_html}
  {exp_html}
  {proj_html}
  {edu_html}
  {cert_html}
    {achievement_html}
</div>"""


# ─── PDF Generation ───────────────────────────────────────────────────────────

# Minimal CSS used when rendering the HTML preview to PDF with PyMuPDF.
# PyMuPDF understands plain CSS well; SCSS features are avoided here.
_PDF_CSS = """
html, body { margin: 0; padding: 0; }
body { font-family: 'Times New Roman', Times, serif; font-size: 13px; line-height: 160%; color: #000000; background: #ffffff; }
.t1-resume { font-family: 'Times New Roman', Times, serif; font-size: 13px; color: #000000; background: white; width: 100%; max-width: 560px; padding: 18px 20px 16px 20px; box-sizing: border-box; margin: 0 auto; }
.t1-logo { text-align: left; margin: 0 0 12px 0; padding: 0; max-width: 100%; }
.t1-logo img { display: block; width: 110px; height: auto; max-width: 100%; }
.t1-name { font-family: 'Inter', sans-serif; font-size: 14px; font-weight: 700; color: #000000; margin: 0 0 4px; text-align: center; }
.t1-contact { font-size: 13px; color: #000000; margin: 0 0 8px; text-align: left; }
.t1-spacer { height: 10px; }
.t1-section-hdr, .t1-proj-num, .t1-role-co, .t1-resp-hdr { font-family: 'Inter', sans-serif; font-size: 14px; font-weight: 700; color: #0072B4; text-transform: uppercase; }
.t1-section-hdr { margin: 12px 0 4px; page-break-after: avoid; break-after: avoid; }
.t1-bullet-list { margin: 0 0 6px 0; padding-left: 20px; list-style-type: disc; list-style-position: outside; }
.t1-bullet { font-size: 13px; margin-bottom: 4px; padding-left: 4px; line-height: 160%; color: #000000; }
.t1-body { font-size: 13px; line-height: 160%; margin: 2px 0; color: #000000; }
.t1-exp-table { width: 100%; border-collapse: collapse; margin-bottom: 8px; }
.t1-exp-table td { padding: 0; vertical-align: top; }
.t1-exp-table tr { page-break-inside: avoid; break-inside: avoid; }
.t1-exp-entry { margin: 0; }
.t1-exp-row { margin: 2px 0; }
.t1-exp-label { font-weight: 600; display: inline-block; width: 130px; }
.t1-exp-sep { display: inline-block; width: 20px; text-align: center; }
.t1-exp-val { display: inline; }
.t1-proj-num { margin: 10px 0 2px; }
.t1-role-co { margin: 0 0 2px; }
.t1-resp-hdr { margin: 4px 0 2px; }

.kf-resume { font-family: 'Times New Roman', Times, serif; font-size: 13px; color: #000000; background: white; width: 100%; max-width: 560px; padding: 18px 20px 16px 20px; box-sizing: border-box; margin: 0 auto; }
.kf-logo { text-align: left; margin: 0 0 12px 0; padding: 0; max-width: 100%; }
.kf-logo img { display: block; width: 110px; height: auto; max-width: 100%; }
.kf-name { font-family: 'Inter', sans-serif; font-size: 14px; font-weight: 700; color: #000000; text-align: center; margin: 0 0 8px; }
.kf-spacer { height: 10px; }
.kf-h1, .kf-proj-num { font-family: 'Inter', sans-serif; font-size: 14px; font-weight: 700; color: #0072B4; text-transform: uppercase; }
.kf-h1 { margin: 10px 0 4px; page-break-after: avoid; break-after: avoid; }
.kf-body { font-size: 13px; margin: 0 0 3px; line-height: 160%; color: #000000; }
.kf-bullet-list { margin: 0 0 6px 0; padding-left: 20px; list-style-type: disc; list-style-position: outside; }
.kf-bullet { font-size: 13px; margin-bottom: 4px; padding-left: 4px; line-height: 160%; color: #000000; }
.kf-tab-row { margin: 2px 0; page-break-inside: avoid; break-inside: avoid; }
.kf-proj-num { margin-top: 8px; }
"""


def convert_html_to_pdf(html_content: str, output_path: str) -> str:
    """Convert HTML resume preview to PDF using PyMuPDF.

    This is the primary PDF path. It avoids Microsoft Word COM automation,
    which is unreliable when run repeatedly from a service.
    """
    full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>{_PDF_CSS}</style>
</head>
<body>
{html_content}
</body>
</html>"""
    # Remove unsupported CSS that can confuse PyMuPDF's HTML engine.
    clean_html = re.sub(r"var\(--[^)]+\)", "#0072B4", full_html)

    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is required for HTML-to-PDF conversion") from exc

    story = fitz.Story(clean_html)
    writer = fitz.DocumentWriter(output_path)
    page_rect = fitz.paper_rect("a4")
    content_rect = page_rect + (20, 20, -20, -20)
    try:
        more = 1
        while more:
            device = writer.begin_page(page_rect)
            more, _ = story.place(content_rect)
            story.draw(device)
            writer.end_page()
    finally:
        writer.close()
    return output_path


# Serialize Word COM PDF conversions. Word on Windows does not tolerate
# multiple concurrent automation sessions well; a lock prevents conflicts.
_pdf_convert_lock = None


def _get_pdf_convert_lock():
    global _pdf_convert_lock
    if _pdf_convert_lock is None:
        import threading
        _pdf_convert_lock = threading.Lock()
    return _pdf_convert_lock


def _wait_for_word_exit(max_wait: int = 30) -> None:
    """Poll until no winword.exe process remains."""
    import subprocess, time
    start = time.time()
    while time.time() - start < max_wait:
        result = subprocess.run(
            ["powershell", "-Command", "(Get-Process winword -ErrorAction SilentlyContinue).Count"],
            capture_output=True, text=True,
        )
        try:
            if int((result.stdout or "0").strip()) == 0:
                return
        except ValueError:
            pass
        time.sleep(1)


def _convert_docx_with_word_com(docx_path: str, pdf_path: str, timeout: int = 120) -> None:
    """Fallback DOCX-to-PDF via Microsoft Word COM automation.

    Kept only as a safety net. Word COM can hang on repeated calls, so the
    HTML-to-PDF path above is preferred.
    """
    import subprocess

    worker_code = r'''
import sys, os, subprocess, time
docx_path = sys.argv[1]
pdf_path = sys.argv[2]

def kill_word():
    subprocess.run(
        ["powershell", "-Command", "Get-Process winword -ErrorAction SilentlyContinue | Stop-Process -Force"],
        capture_output=True,
    )

def wait_for_word_exit(max_wait=30):
    start = time.time()
    while time.time() - start < max_wait:
        result = subprocess.run(
            ["powershell", "-Command", "(Get-Process winword -ErrorAction SilentlyContinue).Count"],
            capture_output=True, text=True,
        )
        try:
            if int((result.stdout or "0").strip()) == 0:
                return
        except ValueError:
            pass
        time.sleep(1)

kill_word()
wait_for_word_exit()

word = None
try:
    import pythoncom
    pythoncom.CoInitialize()
    import win32com.client as wc

    word = wc.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    word.ScreenUpdating = False
    word.Options.ConfirmConversions = False
    word.Options.SavePropertiesPrompt = False
    word.Options.SaveNormalPrompt = False

    doc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True, AddToRecentFiles=False)
    doc.ExportAsFixedFormat(
        OutputFileName=os.path.abspath(pdf_path),
        ExportFormat=17,  # wdExportFormatPDF
        OpenAfterExport=False,
        OptimizeFor=0,    # wdExportOptimizeForPrint
    )
    doc.Close(SaveChanges=False)
    print("OK")
finally:
    if word is not None:
        try:
            word.Quit(SaveChanges=0)
        except Exception:
            pass
    pythoncom.CoUninitialize()
    kill_word()
    wait_for_word_exit()
'''

    proc = subprocess.Popen(
        [sys.executable, "-c", worker_code, docx_path, pdf_path],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    try:
        stdout, stderr = proc.communicate(timeout=timeout)
    except subprocess.TimeoutExpired:
        proc.kill()
        proc.wait()
        _wait_for_word_exit()
        raise RuntimeError("PDF conversion timed out (Word subprocess)")

    _wait_for_word_exit()
    if proc.returncode != 0 or "OK" not in (stdout or "") or not os.path.exists(pdf_path):
        err_msg = (stderr or stdout or "Word PDF conversion failed").strip()
        raise RuntimeError(f"PDF conversion failed: {err_msg}")


def convert_to_pdf(docx_path: str, timeout: int = 120) -> str:
    """Convert docx to pdf. Returns pdf path or raises.

    Tries LibreOffice first, then Word COM automation as a fallback.
    In practice this function is rarely used because PDFs are generated
    eagerly from the HTML preview in _generate_session_artifacts.
    """
    pdf_path = docx_path.replace(".docx", ".pdf")

    # Try LibreOffice first (faster, doesn't require Word)
    try:
        import subprocess
        import shutil
        soffice_paths = [
            shutil.which("soffice"),
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        ]

        for soffice_path in soffice_paths:
            if soffice_path and os.path.exists(soffice_path):
                subprocess.run(
                    [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", str(Path(docx_path).parent), docx_path],
                    check=True,
                    capture_output=True,
                    timeout=60,
                )
                if os.path.exists(pdf_path):
                    return pdf_path
                break
    except Exception:
        pass  # Fall back to Word conversion

    if os.name != "nt":
        raise RuntimeError("PDF conversion requires LibreOffice on non-Windows platforms.")

    with _get_pdf_convert_lock():
        _convert_docx_with_word_com(docx_path, pdf_path, timeout=timeout)
    return pdf_path
