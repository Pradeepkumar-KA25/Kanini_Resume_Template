from docx import Document

def full_analyze(path, label):
    doc = Document(path)
    sec = doc.sections[0]
    print(f"\n====== {label} ======")
    print(f"Page: {sec.page_width.cm:.1f}cm x {sec.page_height.cm:.1f}cm")
    print(f"Margins top={sec.top_margin.cm:.2f} bot={sec.bottom_margin.cm:.2f} left={sec.left_margin.cm:.2f} right={sec.right_margin.cm:.2f}")
    print("\n--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text and i > 3:
            continue
        ri = ""
        for r in p.runs:
            if r.text.strip():
                sz = r.font.size
                col = None
                try:
                    col = str(r.font.color.rgb)
                except Exception:
                    pass
                sz_val = round(sz.pt, 1) if sz else "?"
                ri = f" bold={r.bold} sz={sz_val} col={col}"
                break
        print(f"  [{i:02d}] {p.style.name!r:28s} | {p.text[:100]!r}{ri}")
    print("\n--- TABLES ---")
    for ti, tbl in enumerate(doc.tables):
        print(f"  Table {ti}: {len(tbl.rows)}r x {len(tbl.columns)}c")
        for ri, row in enumerate(tbl.rows[:6]):
            for ci, cell in enumerate(row.cells[:4]):
                bg = None
                try:
                    tcPr = cell._tc.tcPr
                    if tcPr is not None:
                        NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        shd = tcPr.find(f"{{{NS}}}shd")
                        if shd is not None:
                            bg = shd.get(f"{{{NS}}}fill")
                except Exception:
                    pass
                cf = ""
                for cp2 in cell.paragraphs:
                    for cr2 in cp2.runs:
                        if cr2.text.strip():
                            sz2 = cr2.font.size
                            col2 = None
                            try:
                                col2 = str(cr2.font.color.rgb)
                            except Exception:
                                pass
                            sz2_val = round(sz2.pt, 1) if sz2 else "?"
                            cf = f" bold={cr2.bold} sz={sz2_val} col={col2}"
                            break
                    if cf:
                        break
                ct = cell.text[:80].replace("\n", " / ")
                print(f"    [{ri},{ci}] bg={bg}: {ct!r}{cf}")

full_analyze("sample_deloitte2.docx", "DELOITTE FORMAT SAMPLE")
