import asyncio
import json

import fitz
import pytest
from fastapi import HTTPException
from docx import Document

import main
from renderers.render_service import RendererFactory
from renderers.user_template_service import UserTemplateRenderService
from templates.registry import TemplateRegistry
from templates.registry.metadata import TemplateMetadata
from tests.test_template_generation import VALID_SPEC


def _service(tmp_path, spec=VALID_SPEC):
    package = tmp_path / "user-modern"
    package.mkdir()
    (package / "template-spec.json").write_text(json.dumps(spec), encoding="utf-8")
    metadata = TemplateMetadata(id="user-modern", display_name="Modern Resume", description="Test", version="1.0", enabled=True, supported_outputs=["html", "docx", "pdf"], page_size=spec["page"]["size"], user_created=True)
    return UserTemplateRenderService(metadata, package)


def test_user_template_generates_docx_and_pdf(tmp_path, normal_resume):
    service = _service(tmp_path)
    resume = main.ResumeAdapter.from_legacy(normal_resume)
    docx = service.render_docx(resume, tmp_path / "resume.docx").path
    pdf = service.render_pdf(resume, tmp_path / "resume.tex").path
    assert docx and docx.is_file()
    assert pdf and pdf.is_file() and pdf.stat().st_size > 0
    assert "Riya Raman" in "\n".join(paragraph.text for paragraph in Document(docx).paragraphs)


def test_user_template_pdf_handles_long_multi_page_content(tmp_path, normal_resume):
    normal_resume["summary"] = "Long summary content. " * 1500
    service = _service(tmp_path)
    pdf = service.render_pdf(main.ResumeAdapter.from_legacy(normal_resume), tmp_path / "long.tex").path
    document = fitz.open(pdf)
    try:
        assert len(document) > 1
    finally:
        document.close()


def test_user_template_two_column_docx_and_missing_sections(tmp_path, normal_resume):
    spec = {**VALID_SPEC, "layout": {"columns": 2, "sidebar_position": "left", "section_alignment": "left"}}
    normal_resume.update({"summary": "", "projects": [], "certifications": []})
    docx = _service(tmp_path, spec).render_docx(main.ResumeAdapter.from_legacy(normal_resume), tmp_path / "two-column.docx").path
    document = Document(docx)
    assert len(document.tables) == 1
    assert "Projects" not in "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_user_template_download_uses_safe_filename_and_supported_outputs(tmp_path, monkeypatch, normal_resume):
    package = tmp_path / "user-modern"
    package.mkdir()
    manifest = {"id": "user-modern", "display_name": "Modern Resume", "description": "Test", "version": "1.0", "enabled": True, "supported_outputs": ["html", "docx", "pdf"], "page_size": "A4", "aliases": [], "assets": {}, "download_base_name": "user-modern", "user_created": True}
    (package / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    (package / "template-spec.json").write_text(json.dumps(VALID_SPEC), encoding="utf-8")
    registry = TemplateRegistry.discover(user_templates_dir=tmp_path)
    monkeypatch.setattr(main, "TEMPLATE_REGISTRY", registry)
    monkeypatch.setattr(main, "RENDERER_FACTORY", RendererFactory(registry))
    monkeypatch.setattr(main, "TEMP_DIR", tmp_path / "artifacts")
    session_id = "user-output"
    normal_resume["contact"]["name"] = "John Doe / Test"
    main.SESSIONS[session_id] = {"review_data": normal_resume, "resume_data": normal_resume, "files": {}, "filename": "resume.pdf"}
    try:
        response = asyncio.run(main.download_file(session_id, "user-modern", "docx"))
        assert response.media_type.endswith("wordprocessingml.document")
        assert "John_Doe_Test_Modern_Resume.docx" in response.headers["content-disposition"]
        with pytest.raises(HTTPException) as unsupported:
            asyncio.run(main.download_file(session_id, "user-modern", "txt"))
        assert unsupported.value.status_code == 400
        with pytest.raises(HTTPException) as missing:
            asyncio.run(main.download_file(session_id, "unknown", "pdf"))
        assert missing.value.status_code == 400
    finally:
        main.SESSIONS.pop(session_id, None)