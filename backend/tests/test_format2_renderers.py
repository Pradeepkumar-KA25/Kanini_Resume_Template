from pathlib import Path
import shutil

import pymupdf
import pytest
from docx import Document

from models.resume import Project, ResumeData
from renderers.render_service import RendererFactory
from templates.registry import TemplateRegistry


@pytest.fixture
def renderer():
    return RendererFactory(TemplateRegistry.discover()).get("template2")


@pytest.fixture
def resume():
    return ResumeData(contact={"name": "Zoë D’Arcy"}, summary="Data engineer with C# & Python experience.", skills={"Data": ["Python", "SQL", "Spark"]}, experience=[{"company": "Contoso", "title": "Engineer", "dates": "2022 - Present"}], projects=[{"name": "Platform", "client": "Contoso & Fabrikam", "role": "Lead", "description": "Modernized data delivery.", "technologies": ["Python", "SQL"], "responsibilities": ["Delivered 100% coverage."]}], education=[{"degree": "B.Tech", "institution": "Example University", "year": "2020"}])


def test_format2_registry_and_rendered_formats(renderer, resume, tmp_path: Path):
    assert renderer.template_id == "kanini-format-2"
    html = renderer.render_html(resume).content or ""
    docx = renderer.render_docx(resume, tmp_path / "format2.docx")
    tex = renderer.render_latex(resume, tmp_path / "format2.tex")
    assert "@page { size: A4" in html and "Project – I" in html
    assert "WORKING EXPERIENCE:" in "\n".join(p.text for p in Document(docx.path).paragraphs)
    assert r"C\# \& Python" in tex.path.read_text(encoding="utf-8")


def test_format2_includes_all_populated_sections(renderer, tmp_path: Path):
    resume = ResumeData(
        contact={"name": "Candidate Name"},
        summary="Professional summary.",
        skills={"Data": ["Python"]},
        experience=[{"company": "Contoso", "title": "Engineer", "responsibilities": ["Built the experience feature."], "projects": [{"name": "Migration", "description": "Project delivery."}]}],
        certifications=["Azure Certified"],
        achievements=["Award winner"],
    )

    html = renderer.render_html(resume).content or ""
    docx = renderer.render_docx(resume, tmp_path / "format2-all-sections.docx")
    docx_text = "\n".join(paragraph.text for paragraph in Document(docx.path).paragraphs)

    for heading in ("Professional Summary:", "Technical Skills:", "Working Experience:", "Project Summary:", "Certifications:", "Achievements:"):
        assert heading in html
        assert heading.upper() in docx_text
    assert "Built the experience feature." in html
    assert "Built the experience feature." in docx_text


@pytest.mark.skipif(shutil.which("xelatex") is None, reason="XeLaTeX is not installed")
def test_format2_pdf_includes_all_populated_sections(renderer, tmp_path: Path):
    resume = ResumeData(
        contact={"name": "Candidate Name"},
        summary="Professional summary. " * 300,
        skills={"Data": ["Python", "SQL"]},
        experience=[{"company": "Contoso", "title": "Engineer", "responsibilities": ["Delivered a platform. " * 100], "projects": [{"name": "Migration", "description": "Project delivery. " * 100}]}],
        certifications=["Azure Certified"],
        achievements=["Award winner"],
    )

    result = renderer.render_pdf(resume, tmp_path / "format2-all-sections.tex")
    document = pymupdf.open(result.path)
    text = "\n".join(page.get_text() for page in document)

    assert len(document) > 1
    for heading in ("PROFESSIONAL SUMMARY:", "TECHNICAL SKILLS:", "WORKING EXPERIENCE:", "PROJECT SUMMARY:", "CERTIFICATIONS:", "ACHIEVEMENTS:"):
        assert heading in text


@pytest.mark.skipif(shutil.which("xelatex") is None, reason="XeLaTeX is not installed")
@pytest.mark.parametrize("summary", ["Short resume.", "Long content. " * 160, "Zoë works in München.", "Characters: & % $ # _ { } ~ ^ \\ "])
def test_format2_pdf_handles_content_edges(renderer, resume, tmp_path: Path, summary: str):
    resume.summary = summary
    resume.projects.append(Project(name="Second Project", description="Long project. " * 80, responsibilities=["Long responsibility. " * 80]))
    result = renderer.render_pdf(resume, tmp_path / "format2.tex")
    document = pymupdf.open(result.path)
    assert len(document) >= 1
    for page in document:
        assert page.rect.width == pytest.approx(595.28, abs=1)
        assert page.rect.height == pytest.approx(841.89, abs=1)
        assert page.get_image_info(xrefs=True)
        assert all(block[2] <= page.rect.width + 0.1 and block[3] <= page.rect.height + 0.1 for block in page.get_text("blocks"))