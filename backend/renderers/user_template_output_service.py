from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from models.resume import ResumeData
from models.template_spec import TemplateSpec

from .base import RenderResult, RendererError


def render_user_template_pdf(html: str, spec: TemplateSpec, output_path: Path) -> RenderResult:
    """Render the controlled HTML preview to a multi-page PDF with spec page bounds."""
    try:
        import fitz
    except ImportError as exc:
        raise RendererError("PyMuPDF is required for PDF generation.") from exc

    page_rect = fitz.paper_rect("a4" if spec.page.size == "A4" else "letter")
    content_rect = page_rect
    document_html = f"<!doctype html><html><head><meta charset=\"utf-8\"></head><body>{html}</body></html>"
    try:
        story = fitz.Story(document_html)
        writer = fitz.DocumentWriter(str(output_path))
        try:
            more = 1
            while more:
                device = writer.begin_page(page_rect)
                more, _ = story.place(content_rect)
                story.draw(device)
                writer.end_page()
        finally:
            writer.close()
    except Exception as exc:
        raise RendererError("Failed to generate PDF from the user template.") from exc
    if not output_path.is_file() or output_path.stat().st_size == 0:
        raise RendererError("Generated user template PDF is empty.")
    return RenderResult("user-template", "pdf", path=output_path)


def render_user_template_docx(resume: ResumeData, spec: TemplateSpec, output_path: Path) -> RenderResult:
    """Create a Word document directly from structured resume data and a validated spec."""
    document = Document()
    section = document.sections[0]
    _apply_page_spec(section, spec)
    _configure_normal_style(document, spec)
    _add_header(document, resume, spec)

    rendered = [(name, _section_lines(name, resume)) for name in spec.sections]
    rendered = [(name, lines) for name, lines in rendered if lines]
    if spec.layout.columns == 2:
        _add_two_column_sections(document, rendered, spec)
    else:
        for name, lines in rendered:
            _add_section(document, name, lines, spec)
    try:
        document.save(output_path)
    except Exception as exc:
        raise RendererError("Failed to generate DOCX from the user template.") from exc
    return RenderResult("user-template", "docx", path=output_path)


def _apply_page_spec(section, spec: TemplateSpec) -> None:
    if spec.page.size == "A4":
        section.page_width, section.page_height = Inches(8.27), Inches(11.69)
    else:
        section.page_width, section.page_height = Inches(8.5), Inches(11)
    margin = Inches(spec.page.margin_inches)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def _configure_normal_style(document: Document, spec: TemplateSpec) -> None:
    style = document.styles["Normal"]
    style.font.name = spec.typography.font_family
    style.font.size = Pt(spec.typography.base_size_pt)


def _add_header(document: Document, resume: ResumeData, spec: TemplateSpec) -> None:
    alignment = WD_ALIGN_PARAGRAPH.CENTER if spec.header.layout == "centered" else WD_ALIGN_PARAGRAPH.LEFT
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(resume.contact.name or "Candidate Name")
    run.bold = True
    run.font.name = spec.typography.font_family
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    contact_values = [value for value in (resume.contact.email, resume.contact.phone, resume.contact.location, resume.contact.linkedin, resume.contact.github) if value]
    if contact_values:
        contact = document.add_paragraph()
        contact.alignment = alignment
        separator = "\n" if spec.header.contact_layout == "stacked" else " | "
        contact.add_run(separator.join(contact_values))
    if spec.header.show_divider:
        divider = document.add_paragraph()
        divider.paragraph_format.space_after = Pt(spec.spacing.section_gap_pt / 2)
        divider.add_run("_" * 70).font.color.rgb = _color(spec.colors.accent)


def _add_two_column_sections(document: Document, rendered: list[tuple[str, list[str]]], spec: TemplateSpec) -> None:
    sidebar_names = {"skills", "education", "certifications", "achievements"}
    sidebar = [(name, lines) for name, lines in rendered if name in sidebar_names]
    main = [(name, lines) for name, lines in rendered if name not in sidebar_names]
    if not sidebar or not main:
        for name, lines in rendered:
            _add_section(document, name, lines, spec)
        return
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    left, right = table.rows[0].cells
    sidebar_cell, main_cell = (left, right) if spec.layout.sidebar_position == "left" else (right, left)
    _add_sections_to_cell(sidebar_cell, sidebar, spec)
    _add_sections_to_cell(main_cell, main, spec)


def _add_sections_to_cell(cell, rendered: list[tuple[str, list[str]]], spec: TemplateSpec) -> None:
    first = True
    for name, lines in rendered:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        _write_section(paragraph, cell, name, lines, spec)


def _add_section(document: Document, name: str, lines: list[str], spec: TemplateSpec) -> None:
    paragraph = document.add_paragraph()
    _write_section(paragraph, document, name, lines, spec)


def _write_section(paragraph, container, name: str, lines: list[str], spec: TemplateSpec) -> None:
    paragraph.paragraph_format.space_before = Pt(spec.spacing.section_gap_pt)
    heading = paragraph.add_run(name.replace("_", " ").title())
    heading.bold = True
    heading.font.name = spec.typography.font_family
    heading.font.size = Pt(12)
    heading.font.color.rgb = RGBColor(0, 0, 0)
    for line in lines:
        body = container.add_paragraph()
        body.paragraph_format.space_after = Pt(max(2, spec.spacing.section_gap_pt / 3))
        body.add_run(line)


def _section_lines(name: str, resume: ResumeData) -> list[str]:
    if name == "summary": return [line for line in resume.summary.splitlines() if line]
    if name == "skills": return [f"{category}: {', '.join(items)}" for category, items in resume.skills.items() if items]
    if name == "experience": return [_entry_line(item.title, item.company_name or item.company, item.dates, item.location, item.responsibilities) for item in resume.experience]
    if name == "projects": return [_entry_line(item.name, item.client, item.role, item.duration, [item.description, *item.responsibilities]) for item in resume.projects]
    if name == "education": return [" | ".join(part for part in (item.degree, item.institution, item.year, item.gpa) if part) for item in resume.education]
    if name == "certifications": return list(resume.certifications)
    if name == "achievements": return list(resume.achievements)
    return []


def _entry_line(*values) -> str:
    main_values, details = values[:-1], values[-1]
    prefix = " | ".join(str(value) for value in main_values if value)
    detail_values = [str(value) for value in details if value]
    return "\n".join([prefix, *detail_values]) if prefix else "\n".join(detail_values)


def _color(value: str) -> RGBColor:
    return RGBColor.from_string(value.lstrip("#"))