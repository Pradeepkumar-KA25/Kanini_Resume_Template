from docx import Document
from docx.shared import Pt

def analyze_doc(path, label):
    doc = Document(path)
    print(f"\n\n====== {label} ======")
    sec = doc.sections[0]
    print(f"Page: {sec.page_width.cm:.1f}cm x {sec.page_height.cm:.1f}cm")
    print(f"Margins: top={sec.top_margin.cm:.1f} bottom={sec.bottom_margin.cm:.1f} left={sec.left_margin.cm:.1f} right={sec.right_margin.cm:.1f}")

    print("\n--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs[:50]):
        style = p.style.name
        text = p.text[:90] if p.text else "(empty)"
        align = str(p.alignment)
        run_info = ""
        for r in p.runs:
            if r.text.strip():
                sz = r.font.size
                col = None
                try:
                    col = str(r.font.color.rgb)
                except Exception:
                    pass
                run_info = f" bold={r.bold} italic={r.italic} size={round(sz.pt,1) if sz else '?'} color={col}"
                break
        sb = p.paragraph_format.space_before
        sa = p.paragraph_format.space_after
        li = p.paragraph_format.left_indent
        print(f"  [{i:02d}] {style!r:28s} | {text!r}{run_info}")

    print("\n--- TABLES ---")
    for ti, tbl in enumerate(doc.tables[:8]):
        print(f"  Table {ti}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
        for ri, row in enumerate(tbl.rows[:4]):
            for ci, cell in enumerate(row.cells[:4]):
                ct = cell.text[:60].replace("\n", " / ")
                bg = None
                try:
                    tcPr = cell._tc.tcPr
                    if tcPr is not None:
                        shd = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                        if shd is not None:
                            bg = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                except Exception:
                    pass
                print(f"    [{ri},{ci}] bg={bg}: {ct!r}")

analyze_doc("sample_kanini.docx", "SAMPLE KANINI PROFILE")
analyze_doc("sample_kanini2.docx", "Kanini Format  (Downloads)")
