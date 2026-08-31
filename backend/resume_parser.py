import re
import os
import tempfile
import logging
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger("kanini.resume")


def _normalise_pdf_text_layout(text: str) -> str:
    """Repair common PDF extraction layout collapse before section parsing.

    Some resumes arrive as one long line with headings glued together, e.g.
    `PROFILE SUMMARYEDUCATIONSKILLS ... WORK EXPERIENCE ... CONTACT...`.
    This restores boundaries so downstream section parsing can work.
    """
    if not text:
        return ""

    normalized = text.replace("\r\n", "\n").replace("\r", "\n")

    # Recover common mojibake produced by some PDF extraction/terminal encodings.
    normalized = (
        normalized
        .replace("â€“", "-")
        .replace("â€”", "-")
        .replace("âˆ’", "-")
        .replace("â€™", "'")
        .replace("â€œ", '"')
        .replace("â€\x9d", '"')
        .replace("â€¢", "•")
        .replace("â—", "•")
    )

    # Collapse letter-spaced words like "F u l l S t a c k" when they are made of
    # many isolated single-letter tokens.
    normalized = re.sub(
        r"\b(?:[A-Za-z]\s+){3,}[A-Za-z]\b",
        lambda m: m.group(0).replace(" ", ""),
        normalized,
    )

    headings = [
        "PROFILE SUMMARY",
        "PROFESSIONAL SUMMARY",
        "WORK EXPERIENCE",
        "WORKING EXPERIENCE",
        "EMPLOYMENT HISTORY",
        "PROJECT EXPERIENCE",
        "PROJECT SUMMARY",
        "EDUCATION",
        "SKILLS",
        "TECHNICAL SKILLS",
        "CERTIFICATES",
        "CERTIFICATIONS",
        "CONTACT",
        "CORE STRENGTHS",
    ]

    for heading in sorted(headings, key=len, reverse=True):
        normalized = re.sub(
            rf"\s*{re.escape(heading)}\s*",
            f"\n{heading}\n",
            normalized,
            flags=re.IGNORECASE,
        )

    # Contact values often get glued to the CONTACT heading.
    normalized = re.sub(r"\nCONTACT(?=\+|\d|[A-Za-z0-9._%+-]+@)", "\nCONTACT\n", normalized, flags=re.IGNORECASE)

    # Restore likely line breaks before email / linkedin / github / phone when
    # they get appended to preceding text.
    normalized = re.sub(r"(?<!\n)([\w.+-]+@[\w-]+\.[A-Za-z]{2,})", r"\n\1", normalized)
    normalized = re.sub(r"(?<!\n)((?:https?://)?(?:www\.)?linkedin\.com/\S+)", r"\n\1", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"(?<!\n)((?:https?://)?(?:www\.)?github\.com/\S+)", r"\n\1", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"(?<!\n)(\+?\d[\d\s()-]{8,}\d)", r"\n\1", normalized)

    # Normalize excessive blank lines introduced by heading repair.
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()

# ─────────────────────────────────────────────────────────────────────────────
# Text extraction helpers
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_tables_text(file_path: str) -> str:
    """Extract text from PDF tables using pdfplumber as a fallback."""
    try:
        import pdfplumber  # type: ignore

        parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if not tables:
                    continue
                for table in tables:
                    for row in table:
                        row_text = " ".join(str(cell or "") for cell in row if cell)
                        if row_text.strip():
                            parts.append(row_text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_pdf_ocr_text(file_path: str) -> str:
    """OCR fallback for image-based / scanned PDFs using PyMuPDF + pytesseract."""
    try:
        try:
            import pymupdf as fitz  # type: ignore
        except Exception:
            import fitz  # type: ignore
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore

        parts = []
        doc = fitz.open(file_path)
        try:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img)
                if text.strip():
                    parts.append(text)
        finally:
            doc.close()
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    def _score_extracted_pdf_text(text: str) -> int:
        if not text:
            return 0
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if not lines:
            return 0

        score = len(lines)

        # Reward presence of key resume headings and field labels.
        heading_hits = sum(
            1
            for ln in lines
            if re.search(
                r"\b(?:technical\s+skills|work(?:ing)?\s+experience|professional\s+summary"
                r"|company\s+name|designation|duration|education|projects?)\b",
                ln,
                re.IGNORECASE,
            )
        )
        score += heading_hits * 5

        # Penalize extraction that collapses into very long noisy lines.
        long_lines = sum(1 for ln in lines if len(ln) > 220)
        score -= long_lines * 4
        return score

    candidates: List[Tuple[str, str]] = []
    pdf_debug_enabled = os.getenv("PDF_PARSE_DEBUG", "0").strip().lower() in {"1", "true", "yes", "on"}

    # Primary: PyMuPDF often preserves block reading order better for CV tables.
    try:
        try:
            import pymupdf as fitz  # type: ignore
        except Exception:
            import fitz  # type: ignore

        chunks = []
        doc = fitz.open(file_path)
        try:
            for page in doc:
                # "blocks" preserves grouped text regions; sort=True keeps reading order.
                blocks = page.get_text("blocks", sort=True) or []
                if blocks:
                    block_lines = []
                    for blk in blocks:
                        # Tuple format: (x0, y0, x1, y1, text, block_no, block_type)
                        txt = str(blk[4] or "").strip() if len(blk) >= 5 else ""
                        if txt:
                            block_lines.append(txt)
                    if block_lines:
                        chunks.append("\n".join(block_lines))
                        continue

                # Fallback per-page text if blocks are empty.
                txt = page.get_text("text", sort=True) or ""
                if txt.strip():
                    chunks.append(txt)
        finally:
            doc.close()

        fitz_text = "\n".join(chunks).strip()
        if fitz_text:
            candidates.append(("pymupdf", fitz_text))
    except Exception:
        pass

    # Primary: pdfplumber tends to preserve row/label structure better for tabular resumes.
    try:
        import pdfplumber  # type: ignore

        chunks = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # layout=False keeps simpler line flow and usually cleaner label-value rows.
                txt = page.extract_text(layout=False, x_tolerance=2, y_tolerance=3) or ""
                if txt.strip():
                    chunks.append(txt)
        plumber_text = "\n".join(chunks).strip()
        if plumber_text:
            candidates.append(("pdfplumber", plumber_text))
    except Exception:
        pass

    # Fallback: existing pdfminer path.
    try:
        from pdfminer.high_level import extract_text as pdf_extract

        miner_text = (pdf_extract(file_path) or "").strip()
        if miner_text:
            candidates.append(("pdfminer", miner_text))
    except Exception:
        pass

    # Fallbacks for tabular / image-based PDFs when text extraction is weak.
    if not candidates:
        table_text = _extract_pdf_tables_text(file_path)
        if table_text:
            candidates.append(("pdfplumber_tables", table_text))
        ocr_text = _extract_pdf_ocr_text(file_path)
        if ocr_text:
            candidates.append(("ocr", ocr_text))
    else:
        scored = [(name, _score_extracted_pdf_text(text), text) for name, text in candidates]
        best_score = max(s for _, s, _ in scored)
        if best_score < 20:
            table_text = _extract_pdf_tables_text(file_path)
            if table_text:
                candidates.append(("pdfplumber_tables", table_text))
        if best_score < 10:
            ocr_text = _extract_pdf_ocr_text(file_path)
            if ocr_text:
                candidates.append(("ocr", ocr_text))

    if not candidates:
        return ""

    # Pick the best extraction candidate based on structure quality.
    scored: List[Tuple[str, int, str]] = [
        (name, _score_extracted_pdf_text(text), text) for name, text in candidates
    ]
    best_name, best_score, best_text = max(scored, key=lambda x: x[1])

    if pdf_debug_enabled:
        logger.warning(
            "[PDF DEBUG][extract] selected=%s score=%d candidates=%s",
            best_name,
            best_score,
            ", ".join(f"{name}:{score}" for name, score, _ in scored),
        )
        preview = "\\n".join([ln for ln in best_text.splitlines() if ln.strip()][:25])
        logger.warning("[PDF DEBUG][extract] preview=\n%s", preview)

    return _normalise_pdf_text_layout(best_text)


def extract_text_from_docx(file_path: str) -> str:
    def _score_extracted_docx_text(text: str) -> int:
        if not text:
            return 0
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if not lines:
            return 0

        score = len(lines)
        heading_hits = sum(
            1
            for ln in lines
            if re.search(
                r"\b(?:technical\s+skills|work(?:ing)?\s+experience|professional\s+summary"
                r"|education|projects?|company\s+name|designation|duration|email|linkedin)\b",
                ln,
                re.IGNORECASE,
            )
        )
        score += heading_hits * 4
        if re.search(r"@[A-Za-z0-9.-]+", text):
            score += 8
        return score

    from docx import Document
    doc = Document(file_path)
    paragraphs = []

    def _process_para(para):
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            return ""
        is_heading = "Heading" in style or "Title" in style
        if not is_heading and para.runs:
            non_empty = [r for r in para.runs if r.text.strip()]
            if non_empty and all(r.bold for r in non_empty) and len(text) < 60:
                is_heading = True
        return f"__HEADING__ {text}" if is_heading else text

    def _append_from_container(paras, tables):
        for para in paras:
            line = _process_para(para)
            paragraphs.append(line)

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text and text not in paragraphs:
                            paragraphs.insert(0, text)

    # ── Read top-level paragraphs ────────────────────────────────────────────
    for para in doc.paragraphs:
        line = _process_para(para)
        paragraphs.append(line)   # keep blank lines too

    # ── Read table cells (name/contact often lives in a header table) ────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text and text not in paragraphs:
                        paragraphs.insert(0, text)   # table content goes first

    # ── Read headers and footers (many resumes place identity/contact there) ──
    for section in doc.sections:
        try:
            _append_from_container(section.header.paragraphs, section.header.tables)
        except Exception:
            pass
        try:
            _append_from_container(section.footer.paragraphs, section.footer.tables)
        except Exception:
            pass

    docx_text = "\n".join(paragraphs)

    # docx2txt often captures text boxes and alternate layout surfaces that
    # python-docx misses. Prefer the richer extraction.
    candidates = [docx_text]
    try:
        import docx2txt

        extracted = docx2txt.process(file_path) or ""
        if extracted.strip():
            candidates.append(extracted)
    except Exception:
        pass

    best = max(candidates, key=_score_extracted_docx_text)

    # Fallback: explicitly render table rows as lines if normal extraction is weak.
    if _score_extracted_docx_text(best) < 10:
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    str(cell.text or "").strip()
                    for cell in row.cells
                    if str(cell.text or "").strip()
                )
                if row_text:
                    table_lines.append(row_text)
        if table_lines:
            table_candidate = "\n".join(table_lines)
            if _score_extracted_docx_text(table_candidate) > _score_extracted_docx_text(best):
                best = table_candidate

    return best or docx_text


def extract_text_from_doc(file_path: str) -> str:
    try:
        import docx2txt
        text = docx2txt.process(file_path) or ""
        if text.strip():
            return text
    except Exception:
        pass

    # Legacy .doc files often need Microsoft Word to convert them to DOCX first.
    # When Word automation is available on Windows, convert to a temporary DOCX
    # and parse the converted file so the logical structure is preserved.
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        temp_docx = None
        doc = None
        try:
            fd, temp_docx = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=True)
            doc.SaveAs2(temp_docx, FileFormat=12)
            doc.Close(False)
            word.Quit()

            return extract_text_from_docx(temp_docx)
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                word.Quit()
            except Exception:
                pass
            if temp_docx and os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                except Exception:
                    pass
    except Exception:
        return ""


def extract_text(file_path: str, file_type: str) -> str:
    ext = file_type.lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif ext == "docx":
        return extract_text_from_docx(file_path)
    elif ext == "doc":
        return extract_text_from_doc(file_path)
    return ""


# ─────────────────────────────────────────────────────────────────────────────
# Section detection
# ─────────────────────────────────────────────────────────────────────────────

SECTION_KEYWORDS: Dict[str, List[str]] = {
    "summary": [
        "PROFESSIONAL SUMMARY", "SUMMARY", "PROFILE", "ABOUT ME",
        "CAREER OBJECTIVE", "OBJECTIVE", "PROFESSIONAL PROFILE",
        "CAREER SUMMARY", "EXECUTIVE SUMMARY", "PERSONAL STATEMENT",
        "ABOUT", "OVERVIEW", "INTRODUCTION", "PROFESSIONAL OVERVIEW",
        "SUMMARY OF EXPERIENCE", "PROFILE SUMMARY", "CAREER PROFILE",
        "PROFESSIONAL HIGHLIGHTS", "PROFILE HIGHLIGHTS", "CAREER HIGHLIGHTS",
        "PROFILE SNAPSHOT", "CAREER SNAPSHOT", "SUMMARY PROFILE",
    ],
    "experience": [
        "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "WORKING EXPERIENCE", "EXPERIENCE",
        "EMPLOYMENT HISTORY", "WORK HISTORY", "CAREER HISTORY",
        "RELEVANT EXPERIENCE", "PROFESSIONAL BACKGROUND",
        "EMPLOYMENT", "CAREER", "WORK",
        "EXPERIENCE & ACHIEVEMENTS", "COMPANY EXPERIENCE", "INDUSTRY EXPERIENCE",
    ],
    "education": [
        "EDUCATION", "EDUCATIONAL QUALIFICATION", "EDUCATIONAL QUALIFICATIONS", "ACADEMIC BACKGROUND",
        "ACADEMIC QUALIFICATIONS", "ACADEMIC DETAILS", "ACADEMICS",
        "EDUCATIONAL BACKGROUND", "ACADEMIC HISTORY",
        "QUALIFICATIONS", "TRAINING AND EDUCATION", "EDUCATIONAL DETAILS",
    ],
    "skills": [
        "TECHNICAL SKILLS", "CORE COMPETENCIES", "KEY SKILLS", "SKILLS",
        "COMPETENCIES", "TECHNOLOGIES", "TECHNICAL EXPERTISE",
        "SKILLS AND EXPERTISE", "TECHNICAL PROFICIENCIES",
        "TECHNICAL COMPETENCIES", "CORE SKILLS",
        "AREAS OF EXPERTISE", "EXPERTISE", "SKILL SET",
        "TECHNOLOGIES AND TOOLS", "TOOLS AND TECHNOLOGIES",
        "SKILLS & EXPERTISE", "SKILLS & TECHNOLOGIES",
        "TECHNICAL ABILITIES", "IT SKILLS", "SOFTWARE SKILLS",
        "TOOLS", "PROGRAMMING SKILLS", "TECHNICAL STACK",
    ],
    "certifications": [
        "CERTIFICATIONS", "CERTIFICATES", "PROFESSIONAL CERTIFICATIONS",
        "CREDENTIALS", "LICENSES", "PROFESSIONAL LICENSES",
        "CERTIFICATIONS AND TRAINING", "TRAINING",
        "CERTIFICATION", "CERTIFICATE", "PROFESSIONAL DEVELOPMENT",
    ],
    "projects": [
        "PROJECTS", "PROJECT SUMMARY", "PROJECT DESCRIPTION", "PROJECT DETAILS",
        "KEY PROJECTS", "NOTABLE PROJECTS", "PROJECT EXPERIENCE",
        "ACADEMIC PROJECTS", "PROFESSIONAL PROJECTS",
        "PROJECT HIGHLIGHTS", "PROJECT WORK", "RELEVANT PROJECTS",
    ],
    "achievements": [
        "ACHIEVEMENTS", "ACCOMPLISHMENTS", "AWARDS", "HONORS & AWARDS",
        "HONORS", "RECOGNITION", "AWARDS AND ACHIEVEMENTS",
        "KEY ACHIEVEMENTS", "NOTABLE ACHIEVEMENTS", "AWARDS & RECOGNITION",
    ],
}

_SUMMARY_HEADING_ALIASES = [
    "PROFILE SUMMARY",
    "PROFESSIONAL SUMMARY",
    "SUMMARY OF EXPERIENCE",
    "PROFESSIONAL PROFILE",
    "CAREER PROFILE",
    "PROFILE HIGHLIGHTS",
    "PROFESSIONAL HIGHLIGHTS",
    "CAREER SUMMARY",
    "EXECUTIVE SUMMARY",
    "CAREER OBJECTIVE",
    "OBJECTIVE",
    "SUMMARY",
    "PROFILE",
]


def _is_summary_candidate_line(line: str) -> bool:
    text = str(line or "").strip()
    if not text:
        return False

    if EMAIL_RE.search(text) or PHONE_RE.search(text) or LINKEDIN_RE.search(text) or GITHUB_RE.search(text):
        return False

    heading_like = bool(_detect_section(text))
    if heading_like:
        return False

    low = text.casefold()
    if re.search(r"\b(?:company\s+name|designation|duration|roles?\s+and\s+responsibilities|project\s+summary)\b", low):
        return False

    # Category-like skill lines are generally not profile summary content.
    if re.match(r"^[A-Za-z][A-Za-z\s&/\-]{1,35}:\s*[^.]{1,120}$", text) and "," in text:
        return False

    bullet_like = bool(re.match(r"^\s*(?:[•\u2022\u2023\u25E6\u2043\u2219\-\*\u2192\u25BA\uf0b7]+|\d+[.)])\s*", text))
    if bullet_like:
        return len(re.sub(r"^\s*(?:[•\u2022\u2023\u25E6\u2043\u2219\-\*\u2192\u25BA\uf0b7]+|\d+[.)])\s*", "", text).split()) >= 4

    # Narrative summary lines are usually sentence-like and not too short.
    return len(text.split()) >= 7


def _extract_summary_from_raw_text(raw_text: str) -> List[str]:
    if not raw_text:
        return []

    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    stop_headings = sorted(
        {
            kw
            for section, kws in SECTION_KEYWORDS.items()
            if section != "summary"
            for kw in kws
            if len(kw.split()) >= 1
        },
        key=len,
        reverse=True,
    )
    summary_heads = sorted(set(_SUMMARY_HEADING_ALIASES), key=len, reverse=True)

    stop_pat = "|".join(re.escape(h) for h in stop_headings)
    summary_pat = "|".join(re.escape(h) for h in summary_heads)

    # Capture text after a summary heading until the next known section heading.
    block_re = re.compile(
        rf"(?is)(?:^|\n)\s*(?:{summary_pat})\s*:?\s*(.+?)(?=\n\s*(?:{stop_pat})\s*:?(?:\n|$)|\Z)"
    )

    m = block_re.search(text)
    if not m:
        return []

    candidate = m.group(1)
    lines = [ln.strip() for ln in candidate.split("\n") if ln.strip()]
    cleaned = [ln for ln in lines if _is_summary_candidate_line(ln)]
    return cleaned[:10]


def _extract_summary_from_header_lines(header_lines: List[str]) -> List[str]:
    if not header_lines:
        return []

    candidates: List[str] = []
    for line in header_lines:
        text = str(line or "").strip()
        if not text:
            continue
        if _is_summary_candidate_line(text):
            candidates.append(text)

    return candidates[:8]


def _detect_section(line: str) -> Optional[str]:
    # Strip heading prefix, trailing colon/whitespace, then upper-case
    cleaned = re.sub(r"__HEADING__\s*", "", line).strip()
    cleaned = re.sub(r"[:：]+$", "", cleaned).strip()   # remove trailing colon
    cleaned = re.sub(r"\s+", " ", cleaned).upper()

    for section, keywords in SECTION_KEYWORDS.items():
        if cleaned in keywords:
            return section
        for kw in keywords:
            if cleaned == kw:
                return section
            # Prefix match: "WORK EXPERIENCE – FULL TIME" still hits "WORK EXPERIENCE"
            if cleaned.startswith(kw) and (
                len(cleaned) == len(kw) or not cleaned[len(kw)].isalpha()
            ):
                return section

    # Fallback: all-caps line ≤ 6 words, no digits → fuzzy keyword search
    # Only use MULTI-WORD keywords to avoid false-positive matches on company names
    if (
        re.match(r"^[A-Z][A-Z\s&/\-]{2,50}$", cleaned)
        and len(cleaned.split()) <= 6
        and not re.search(r"\d", cleaned)
    ):
        for section, keywords in SECTION_KEYWORDS.items():
            for kw in keywords:
                # Skip single-word keywords — they are already handled above
                if len(kw.split()) < 2:
                    continue
                if kw in cleaned:
                    return section
    return None


def split_sections(text: str) -> Dict[str, List[str]]:
    sections: Dict[str, List[str]] = {"header": []}
    current = "header"
    
    # Pre-pass: detect all section headers and their line positions
    section_positions = []
    for line_num, line in enumerate(text.splitlines()):
        stripped = line.strip()
        detected = _detect_section(stripped)
        if detected:
            section_positions.append((line_num, detected, stripped))

    for line in text.splitlines():
        raw = line.rstrip()
        stripped = raw.strip()
        if not stripped:
            sections.setdefault(current, []).append("")
            continue

        normalized = re.sub(r"^__HEADING__\s*", "", stripped).strip().rstrip(":").upper()
        if normalized in {"DESCRIPTION OF PROJECT", "ROLES AND RESPONSIBILITIES", "RESPONSIBILITIES"}:
            current = "projects"
            sections.setdefault(current, []).append(normalized)
            continue

        detected = _detect_section(stripped)
        if detected:
            current = detected
            sections.setdefault(current, [])
            # Don't add the section header itself to the section content
        else:
            # Strip __HEADING__ prefix so heading-styled lines (company names,
            # titles, etc.) don't bleed into parsed fields with the marker text.
            clean = re.sub(r"^__HEADING__\s*", "", raw)
            sections.setdefault(current, []).append(clean)

    # Clean up: remove trailing empty strings from sections
    for section_name in sections:
        while sections[section_name] and sections[section_name][-1].strip() == "":
            sections[section_name].pop()
    
    return sections


# ─────────────────────────────────────────────────────────────────────────────
# Contact info extraction
# ─────────────────────────────────────────────────────────────────────────────

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?(?:\(?\d{3}\)?[\s\-.]?)?\d{3}[\s\-.]?\d{4}"
)
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+/?", re.IGNORECASE)
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+/?", re.IGNORECASE)


def extract_contact_info(header_lines: List[str]) -> Dict:
    text = "\n".join(header_lines)
    email = (EMAIL_RE.search(text) or type("", (), {"group": lambda *_: ""})()).group(0)
    phone_match = PHONE_RE.search(re.sub(r"\d{4}[-–]\d{4}", "", text))  # avoid year ranges
    phone = phone_match.group(0).strip() if phone_match else ""
    linkedin = (LINKEDIN_RE.search(text) or type("", (), {"group": lambda *_: ""})()).group(0)
    github = (GITHUB_RE.search(text) or type("", (), {"group": lambda *_: ""})()).group(0)

    # Name: first short non-contact line at the top (2–5 words, mostly letters)
    name = ""
    _NO_CONTACT = lambda s: (
        not EMAIL_RE.search(s)
        and not PHONE_RE.search(s)
        and not LINKEDIN_RE.search(s)
        and not re.search(r"https?://|www\.", s, re.I)
    )
    for line in header_lines:
        l = re.sub(r"__HEADING__\s*", "", line).strip()
        if not l:
            continue
        if not _NO_CONTACT(l):
            continue
        words = l.split()
        # A name: 2–6 words, each word starts with a letter, overall ≤ 70 chars
        # Accept ALL-CAPS ("INDIRA ESWARAN") or Title Case ("Indira Eswaran")
        if 2 <= len(words) <= 6 and len(l) <= 70:
            if re.match(r"^[A-Za-z]", l) and re.match(r"^[A-Za-z\s.\-']+$", l):
                # Reject lines that look like job titles or cities with digits
                if not re.search(r"\d", l):
                    name = l.title()   # normalise to Title Case
                    break

    # Location: look for city/state patterns
    location = ""
    loc_patterns = [
        re.compile(r"\b[A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*[A-Z]{2}\b"),  # City, ST
        re.compile(r"\b[A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*[A-Z][a-z]+\b"),  # City, Country
    ]
    for pat in loc_patterns:
        m = pat.search(text)
        if m:
            location = m.group(0)
            break

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "location": location,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Skills extraction
# ─────────────────────────────────────────────────────────────────────────────

TECH_SKILLS_POOL = [
    # Languages
    "Python", "Java", "JavaScript", "TypeScript", "C#", "C++", "C", "Go",
    "Rust", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R", "MATLAB",
    # Web
    "HTML", "CSS", "React", "Angular", "Vue", "Next.js", "Node.js", "Express",
    "Django", "Flask", "FastAPI", "Spring Boot", "ASP.NET", "Laravel",
    # Data / AI
    "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Keras",
    "scikit-learn", "Pandas", "NumPy", "NLP", "Computer Vision",
    # Cloud
    "AWS", "Azure", "GCP", "Google Cloud", "Kubernetes", "Docker",
    "Terraform", "Ansible", "Jenkins", "CI/CD",
    # Databases
    "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch", "Oracle",
    "SQL Server", "SQLite", "DynamoDB", "Cassandra",
    # Tools
    "Git", "GitHub", "GitLab", "Bitbucket", "JIRA", "Confluence",
    "Postman", "Swagger", "GraphQL", "REST", "gRPC", "Microservices",
    # Concepts
    "Agile", "Scrum", "DevOps", "SDLC", "OOP", "Design Patterns",
    "Data Structures", "Algorithms",
]

SKILL_CATEGORIES = {
    "Programming Languages": [
        "Python", "Java", "JavaScript", "TypeScript", "C#", "C++", "C",
        "Go", "Rust", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R", "MATLAB",
    ],
    "Web & Frameworks": [
        "HTML", "CSS", "React", "Angular", "Vue", "Next.js", "Node.js",
        "Express", "Django", "Flask", "FastAPI", "Spring Boot", "ASP.NET",
        "Laravel",
    ],
    "Cloud & DevOps": [
        "AWS", "Azure", "GCP", "Google Cloud", "Kubernetes", "Docker",
        "Terraform", "Ansible", "Jenkins", "CI/CD",
    ],
    "Databases": [
        "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch", "Oracle",
        "SQL Server", "SQLite", "DynamoDB", "Cassandra",
    ],
    "AI & Data Science": [
        "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Keras",
        "scikit-learn", "Pandas", "NumPy", "NLP", "Computer Vision",
    ],
    "Tools & Practices": [
        "Git", "GitHub", "GitLab", "Bitbucket", "JIRA", "Confluence",
        "Postman", "Swagger", "GraphQL", "REST", "Agile", "Scrum", "DevOps",
        "Microservices", "Design Patterns",
    ],
}


def parse_skills(lines: List[str]) -> Dict[str, List[str]]:
    # Ignore lines that are likely from experience/employment blocks.
    exp_noise_re = re.compile(
        r"\b(?:company(?:\s+name)?|designation|duration|employer|organization|organisation"
        r"|responsibilities?|roles?\s+and\s+responsibilities?|client|project\s+[ivx]+|project\s+\d+)\b",
        re.IGNORECASE,
    )
    exp_date_re = ENTRY_DATE_RE
    # Aggressive filters for experience-like content:
    # - Lines with "designed", "developed", "built", "implemented", "led" (typical job responsibility verbs)
    # - Lines with location patterns (City, Country/State)
    # - Lines that look like company names (e.g., "Acme Inc | Role | Location")
    exp_verb_re = re.compile(
        r"\b(?:designed|developed|built|implemented|led|managed|created|deployed|architected"
        r"|migrated|optimized|enhanced|improved|integrated|coordinated|reduced|increased|achieved)\b",
        re.IGNORECASE,
    )
    job_format_re = re.compile(
        r"^[A-Z][A-Za-z\s&]+(?:\s*\|\s*)?[A-Z][A-Za-z\s]+(?:\s*\|\s*)?[A-Z][A-Za-z,\s]+$"
    )

    filtered_lines: List[str] = []
    for ln in lines:
        txt = str(ln or "").strip()
        if not txt:
            continue
        if exp_noise_re.search(txt):
            continue
        if exp_date_re.search(txt):
            continue
        # Skip lines that start with bullet points and contain action verbs (job responsibilities)
        if BULLET_RE.match(txt) and exp_verb_re.search(txt):
            continue
        # Skip lines that look like "Company | Role | Location" format
        if job_format_re.match(txt) and "|" in txt:
            continue
        filtered_lines.append(txt)

    full_text = " ".join(filtered_lines)
    found: List[str] = []

    for skill in TECH_SKILLS_POOL:
        pattern = re.compile(r"\b" + re.escape(skill) + r"\b", re.IGNORECASE)
        if pattern.search(full_text):
            found.append(skill)

    # Also extract bullet-listed / comma-separated skills
    for line in filtered_lines:
        line = re.sub(r"^[\•\-\*\→\►]\s*", "", line.strip())
        # If line is "Category Label: skill1, skill2" only take the values part
        colon_match = re.match(r'^[A-Za-z][A-Za-z\s&/\-]+:\s*(.+)$', line)
        if colon_match:
            line = colon_match.group(1)
        parts = re.split(r"[,|;]", line)  # no / split to preserve CI/CD
        for part in parts:
            skill = part.strip().rstrip(".")
            if 2 <= len(skill) <= 40 and skill not in found:
                found.append(skill)

    # Categorise
    categorised: Dict[str, List[str]] = {}
    uncategorised: List[str] = []

    for skill in found:
        placed = False
        for cat, cat_skills in SKILL_CATEGORIES.items():
            if any(skill.lower() == s.lower() for s in cat_skills):
                categorised.setdefault(cat, []).append(skill)
                placed = True
                break
        if not placed:
            uncategorised.append(skill)

    if uncategorised:
        categorised["Other Skills"] = uncategorised

    return categorised


# ─────────────────────────────────────────────────────────────────────────────
# Experience parsing
# ─────────────────────────────────────────────────────────────────────────────

DATE_RE = re.compile(
    r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?"
    r")\s*(?:\'|'|`)?\d{2,4}|(?:\d{1,2}[\/\-]\d{4}|\d{4})",
    re.IGNORECASE,
)
# Flexible date pattern: range OR single month-year / year
ENTRY_DATE_RE = re.compile(
    r"(?:"
    # Month Year – Month Year / Present / Current
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"[.,\s]+\d{2,4}"
    r"(?:\s*[-–—/to]+\s*"
    r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"[.,\s]+\d{2,4}|Present|Current|Till\s+Date|Today|\d{4}))?"
    r"|"
    # YYYY – YYYY / Present
    r"\d{4}\s*[-–—to]+\s*(?:\d{4}|Present|Current|Till\s+Date|Today)"
    r"|"
    # MM/YYYY – MM/YYYY / Present
    r"\d{1,2}/\d{4}\s*[-–—]+\s*(?:\d{1,2}/\d{4}|Present|Current)"
    r")",
    re.IGNORECASE,
)

BULLET_RE = re.compile(r"^[\s]*[•\-\*\→\►\–]")

EXP_LABEL_RE = re.compile(
    r"^(Company Name|Designation|Duration|Role|Position|Client|Employer|"
    r"Organization|Organisation|Company|Location)\s*[:\-]?\s*(.*)$",
    re.IGNORECASE,
)


def _split_bullet_items(text: str) -> List[str]:
    parts = re.split(r"\s*[•\u2022\u2023\u25E6\u2043\u2219\uf0b7\-\*\u2192\u25BA]+\s+", str(text or "").strip())
    return [p.strip() for p in parts if p and p.strip()]


def _coalesce_exp_label_lines(lines: List[str]) -> List[str]:
    """Join label/value pairs split across multiple lines from PDF/Word extraction.

    Example recovered shape:
      Company Name\n:\nExterro R&D Pvt Ltd.
    becomes:
      Company Name: Exterro R&D Pvt Ltd.
    """
    out: List[str] = []
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        line = (raw or "").strip()
        if not line:
            out.append("")
            i += 1
            continue

        # Standalone separator lines are usually part of split table rows.
        if line in {":", "-", "–", "—"}:
            i += 1
            continue

        m = EXP_LABEL_RE.match(line)
        if not m:
            out.append(line)
            i += 1
            continue

        label = m.group(1).strip()
        value = m.group(2).strip()

        # If value is already on the same line, normalize and keep.
        if value:
            out.append(f"{label}: {value}")
            i += 1
            continue

        # Otherwise pull the next meaningful line as the value, skipping separators.
        j = i + 1
        while j < n and not (lines[j] or "").strip():
            j += 1
        if j < n and (lines[j] or "").strip() in {":", "-", "–", "—"}:
            j += 1
            while j < n and not (lines[j] or "").strip():
                j += 1

        if j < n:
            candidate = (lines[j] or "").strip()
            candidate_m = EXP_LABEL_RE.match(candidate)
            if candidate and not candidate_m:
                out.append(f"{label}: {candidate}")
                i = j + 1
                continue

        # Label-only line with no visible value.
        out.append(f"{label}:")
        i += 1

    return out


def _split_exp_blocks(lines: List[str]) -> List[List[str]]:
    """Group experience lines into per-job blocks.
    A new block starts on a blank line OR when a date line follows bullet lines.
    """
    blocks: List[List[str]] = []
    current: List[str] = []
    prev_had_bullet = False
    current_has_date = False

    company_or_role_re = re.compile(
        r"^(?:Company(?:\s+Name)?|Client|Employer|Organization|Organisation|"
        r"Designation|Role|Duration|Location)\s*[:\-]",
        re.IGNORECASE,
    )

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                blocks.append(current)
                current = []
            prev_had_bullet = False
            current_has_date = False
            continue

        is_bullet = bool(BULLET_RE.match(stripped))
        has_date  = bool(ENTRY_DATE_RE.search(stripped))
        has_label = bool(company_or_role_re.match(stripped))

        # If we hit a date line that isn't a bullet AND previous block already
        # contained bullets → the date belongs to a NEW entry
        if has_date and not is_bullet and prev_had_bullet and current:
            blocks.append(current)
            current = []
            prev_had_bullet = False
            current_has_date = False

        # If a new labeled company/role line starts after a complete date-bearing
        # block, split into a new experience entry.
        if has_label and current and current_has_date and not is_bullet and not has_date:
            blocks.append(current)
            current = []
            prev_had_bullet = False
            current_has_date = False

        current.append(stripped)
        if is_bullet:
            prev_had_bullet = True
        if has_date:
            current_has_date = True

    if current:
        blocks.append(current)
    return blocks


def _parse_exp_block(block: List[str]) -> Optional[Dict]:
    """Extract a single experience entry from a block of lines."""
    date_str = ""
    header_lines: List[str] = []   # non-bullet, non-date text lines
    responsibilities: List[str] = []
    title_field = ""
    company_field = ""
    location_field = ""
    in_responsibilities = False

    def _strip_field_label(text: str) -> Tuple[str, str]:
        m = re.match(
            r"^(Company Name|Designation|Duration|Role|Client|Employer|"
            r"Organization|Organisation|Company|Location|Position|"
            r"Roles and Responsibilities|Responsibilities)\s*[:\-]\s*(.*)$",
            text,
            re.IGNORECASE,
        )
        if m:
            return m.group(1).strip().lower(), m.group(2).strip()
        return "", text

    for line in block:
        is_bullet = bool(BULLET_RE.match(line))
        date_match = ENTRY_DATE_RE.search(line)

        label, cleaned_line = _strip_field_label(line)
        if label in {"roles and responsibilities", "responsibilities"}:
            in_responsibilities = True
            if cleaned_line:
                responsibilities.extend(_split_bullet_items(cleaned_line))
            continue
        if label == "duration" and cleaned_line:
            if not date_str:
                date_str = cleaned_line
            continue
        if label in {"company name", "client", "company", "employer", "organization", "organisation"} and cleaned_line:
            company_field = cleaned_line
            in_responsibilities = False
            continue
        if label in {"designation", "role", "position"} and cleaned_line:
            title_field = cleaned_line
            in_responsibilities = False
            continue
        if label == "location" and cleaned_line:
            location_field = cleaned_line
            in_responsibilities = False
            continue

        if is_bullet:
            resp = re.sub(r"^[\s•\-\*\→\►\–]+", "", line).strip()
            if resp:
                responsibilities.extend(_split_bullet_items(resp))
            in_responsibilities = True
        elif date_match:
            if not date_str:
                date_str = date_match.group(0).strip()
            # Keep any text on the same line as the date (title | company)
            rest = line[:date_match.start()].strip(" |–-—()")
            rest2 = line[date_match.end():].strip(" |–-—()")
            for part in [rest, rest2]:
                if part and len(part) > 1:
                    header_lines.append(part)
            in_responsibilities = False
        elif in_responsibilities:
            # Continue collecting responsibility lines even when bullets are absent.
            responsibilities.extend(_split_bullet_items(cleaned_line))
        else:
            header_lines.append(cleaned_line)

    # Build title / company from header lines
    title, company, location = title_field, company_field, location_field
    for part in header_lines:
        # Split on pipe, @, or 3+ spaces — handle 2 or 3 segments
        segs = [s.strip() for s in re.split(r"\s*[|@·]\s*|\s{3,}", part) if s.strip()]
        if len(segs) >= 3:
            if not title and not company:
                title, company = segs[0], segs[1]
                if not location:
                    location = segs[2]
            elif not location:
                location = segs[-1]
        elif len(segs) == 2:
            if not title:
                title, company = segs[0], segs[1]
            elif not company:
                company = segs[0]
        elif segs:
            if not title:
                title = segs[0]
            elif not company:
                company = segs[0]
            elif not location and "," in segs[0] and len(segs[0]) < 50:
                location = segs[0]

    # Common Kanini-style lines may use explicit labels in the extracted text.
    title = re.sub(r"^(Designation|Role|Position|Title)\s*[:\-]\s*", "", title, flags=re.IGNORECASE).strip()
    company = re.sub(r"^(Company Name|Client|Employer|Organization|Organisation|Company)\s*[:\-]\s*", "", company, flags=re.IGNORECASE).strip()
    location = re.sub(r"^(Location)\s*[:\-]\s*", "", location, flags=re.IGNORECASE).strip()

    if not title and not company:
        return None

    # Swap title/company if company field looks like a job role, not an org name
    _ROLE = re.compile(
        r'\b(?:Engineer|Developer|Manager|Analyst|Architect|Designer'
        r'|Consultant|Lead|Senior|Junior|Director|Head|Chief|Officer'
        r'|Programmer|Specialist|Associate|Coordinator|Intern|Trainee)\b',
        re.IGNORECASE,
    )
    if _ROLE.search(company) and not _ROLE.search(title):
        title, company = company, title

    # If responsibilities are still empty, treat long residual lines as duties.
    if not responsibilities:
        for part in header_lines:
            text = part.strip()
            if not text:
                continue
            if ENTRY_DATE_RE.search(text):
                continue
            if re.match(r"^(Company Name|Designation|Duration|Role|Location)\b", text, re.IGNORECASE):
                continue
            if text in {title, company, location}:
                continue
            if len(text.split()) >= 6:
                responsibilities.append(text)

    # Dedupe responsibilities while preserving order.
    deduped_resps: List[str] = []
    seen = set()
    for r in responsibilities:
        rr = str(r).strip()
        if not rr:
            continue
        key = rr.casefold()
        if key in seen:
            continue
        seen.add(key)
        deduped_resps.append(rr)

    return {
        "title": title, "company": company,
        "dates": date_str, "location": location,
        "responsibilities": deduped_resps,
    }


def parse_experience(lines: List[str]) -> List[Dict]:
    lines = _coalesce_exp_label_lines(lines)
    blocks  = _split_exp_blocks(lines)
    entries = [_parse_exp_block(b) for b in blocks]
    return [e for e in entries if e and (e.get("title") or e.get("company"))]


# ─────────────────────────────────────────────────────────────────────────────
# Education parsing
# ─────────────────────────────────────────────────────────────────────────────

DEGREE_KEYWORDS = re.compile(
    r"\b(?:B\.?Tech|M\.?Tech|B\.?E|M\.?E|B\.?Sc|M\.?Sc|B\.?A|M\.?A|Ph\.?D|MBA"
    r"|Bachelor|Master|Associate|Diploma|B\.?Com|M\.?Com)\b",
    re.IGNORECASE,
)


def parse_education(lines: List[str]) -> List[Dict]:
    entries = []
    current: Optional[Dict] = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        year_match = re.search(r"\b(19|20)\d{2}\b", stripped)
        degree_match = DEGREE_KEYWORDS.search(stripped)

        # Handle pipe-separated single-line: "Degree | Institution | Year"
        pipe_segs = [s.strip() for s in stripped.split("|") if s.strip()]
        if len(pipe_segs) >= 2 and (degree_match or year_match):
            year = ""
            deg_part = pipe_segs[0]
            inst_part = pipe_segs[1] if len(pipe_segs) > 1 else ""
            # Year may be in any segment
            for seg in pipe_segs:
                ym = re.search(r"\b(19|20)\d{2}\b", seg)
                if ym:
                    year = ym.group(0)
                    if seg.strip() == year:  # segment is only a year
                        continue
            deg_part = re.sub(r"\b(19|20)\d{2}\b", "", deg_part).strip(" ,-|")
            inst_part = re.sub(r"\b(19|20)\d{2}\b", "", inst_part).strip(" ,-|")
            if current:
                entries.append(current)
            current = {
                "degree": deg_part,
                "institution": inst_part,
                "year": year,
                "gpa": "",
            }
            entries.append(current)
            current = None
            continue

        if degree_match or (year_match and not current):
            if current:
                entries.append(current)
            year = year_match.group(0) if year_match else ""
            current = {
                "degree": stripped.replace(year, "").strip(" ,-"),
                "institution": "",
                "year": year,
                "gpa": "",
            }
        elif current and not current["institution"]:
            current["institution"] = stripped
        elif current:
            gpa_match = re.search(r"GPA[:\s]+(\d+\.?\d*)", stripped, re.IGNORECASE)
            if gpa_match:
                current["gpa"] = gpa_match.group(1)

    if current:
        entries.append(current)

    return [e for e in entries if e.get("degree")]


# ─────────────────────────────────────────────────────────────────────────────
# Certifications / projects / achievements
# ─────────────────────────────────────────────────────────────────────────────

def parse_list_section(lines: List[str]) -> List[str]:
    bullet_split_re = re.compile(r"\s*[•\u2022\u2023\u25E6\u2043\u2219\-\*\u2192\u25BA]+\s+")
    leading_marker_re = re.compile(r"^\s*(?:[•\u2022\u2023\u25E6\u2043\u2219\-\*\u2192\u25BA]+|\d+[.)])\s*")

    items: List[str] = []
    for line in lines:
        raw = line.strip()
        if not raw:
            continue

        has_bullet_delimiter = bool(bullet_split_re.search(raw))
        if has_bullet_delimiter:
            parts = [p.strip() for p in bullet_split_re.split(raw) if p.strip()]
            if parts:
                items.extend(parts)
                continue

        stripped = leading_marker_re.sub("", raw).strip()
        if stripped:
            items.append(stripped)

    return items


def parse_projects(lines: List[str]) -> List[Dict]:
    projects: List[Dict] = []
    current: Optional[Dict] = None
    mode = "header"

    def _dedupe_keep_order(items: List[str]) -> List[str]:
        seen = set()
        out = []
        for item in items:
            text = str(item).strip()
            if not text:
                continue
            key = text.casefold()
            if key in seen:
                continue
            seen.add(key)
            out.append(text)
        return out

    def _is_project_heading(text: str) -> bool:
        return bool(re.match(r"^Project\s+(?:[IVX]+|\d+)\s*[:\-–—]?$", text.strip(), re.IGNORECASE))

    def _new_project(name: str) -> Dict:
        return {
            "name": name,
            "description": "",
            "technologies": [],
            "client": "",
            "role": "",
            "responsibilities": [],
        }

    def _has_content(proj: Dict) -> bool:
        return bool(
            proj and (proj.get("description") or proj.get("responsibilities")
                      or proj.get("technologies") or proj.get("client") or proj.get("role"))
        )

    def _looks_like_project_title(text: str) -> bool:
        """A short, capitalised, non-sentence line that likely names a new project."""
        t = text.strip()
        if not t or len(t) > 80 or t.endswith((".", ",", ";", ":")):
            return False
        if len(t.split()) > 10:
            return False
        return bool(re.match(r"^[A-Z0-9\"'(]", t))

    def _flush():
        nonlocal current
        if current and current.get("name"):
            projects.append(current)
        current = None

    blank_seen = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            blank_seen = True
            continue

        clean = re.sub(r"^__HEADING__\s*", "", stripped).strip()

        # Explicit "Project I:" heading always starts a new project.
        if _is_project_heading(clean):
            _flush()
            current = _new_project(clean.rstrip(":-–— "))
            mode = "header"
            blank_seen = False
            continue

        if current is None:
            current = _new_project(clean)
            mode = "header"
            blank_seen = False
            continue

        # Strip a single leading bullet for inspection, but remember it was one.
        is_bullet = bool(re.match(r"^[\•\-\*\→\►]", clean))
        body = re.sub(r"^[\•\-\*\→\►]\s*", "", clean).strip() if is_bullet else clean

        # Section-only headers ("Description of Project" / "Responsibilities").
        low = body.rstrip(":").strip().lower()
        if low in {"description of project", "project description",
                   "roles and responsibilities", "responsibilities"}:
            mode = "description" if "description" in low else "responsibilities"
            blank_seen = False
            continue

        # Labelled fields — recognised even when prefixed by a bullet.
        label_match = re.match(
            r"^(Client|Customer|Technologies?|Tech Stack|Tools?|Built with|Using|Environment"
            r"|Role|Description of Project|Project Description"
            r"|Roles and Responsibilities|Responsibilities)\s*[:\-]\s*(.*)$",
            body, re.IGNORECASE,
        )
        if label_match:
            label = label_match.group(1).lower()
            value = label_match.group(2).strip()
            if label in {"client", "customer"}:
                current["client"] = value
                mode = "header"
            elif label in {"technologies", "technology", "tech stack", "tool",
                           "tools", "built with", "using", "environment"}:
                techs = [t.strip() for t in re.split(r"[,|/;]", value) if t.strip()]
                current["technologies"].extend(techs)
                mode = "header"
            elif label == "role":
                current["role"] = value
                mode = "header"
            elif label in {"description of project", "project description"}:
                mode = "description"
                if value:
                    current["description"] = (
                        current["description"] + " " + value).strip() if current["description"] else value
            elif label in {"roles and responsibilities", "responsibilities"}:
                mode = "responsibilities"
                if value:
                    current["responsibilities"].append(value)
            blank_seen = False
            continue

        # A non-bullet, title-like line following a blank line starts a NEW
        # project (resumes list projects by plain title, not "Project I").
        if (not is_bullet and blank_seen and _has_content(current)
                and _looks_like_project_title(body)):
            _flush()
            current = _new_project(body)
            mode = "header"
            blank_seen = False
            continue

        blank_seen = False

        if is_bullet:
            if mode == "responsibilities":
                current["responsibilities"].append(body)
            else:
                current["description"] = (
                    current["description"] + " " + body).strip() if current["description"] else body
            continue

        if mode == "responsibilities":
            current["responsibilities"].append(body)
        elif mode == "description":
            current["description"] = (
                current["description"] + " " + body).strip() if current["description"] else body
        else:
            # Heuristic lines in the project header area
            if not current["client"]:
                current["client"] = body
            elif not current["role"] and not _is_project_heading(body):
                current["role"] = body
            elif not current["description"]:
                current["description"] = body

    _flush()

    # Normalize fields and keep only useful projects
    cleaned_projects: List[Dict] = []
    for p in projects:
        p["technologies"] = _dedupe_keep_order([t for t in p.get("technologies", []) if t])
        p["responsibilities"] = _dedupe_keep_order([r for r in p.get("responsibilities", []) if r])
        if p.get("name") or p.get("client") or p.get("description"):
            cleaned_projects.append({
                "name": p.get("name", ""),
                "client": p.get("client", ""),
                "role": p.get("role", ""),
                "description": p.get("description", ""),
                "technologies": p.get("technologies", []),
                "responsibilities": p.get("responsibilities", []),
            })

    return cleaned_projects


# ─────────────────────────────────────────────────────────────────────────────
# Master parse function
# ─────────────────────────────────────────────────────────────────────────────
# Helpers to rescue name/contact that landed inside the summary section
# (happens when resume starts with a heading like "Profile" instead of a
#  plain header block)
# ─────────────────────────────────────────────────────────────────────────────

def _salvage_contact_from_summary(contact: Dict, summary_lines: List[str]):
    """If no name in header, try to pull name+contact from the top of summary."""
    if contact.get("name"):
        return contact, summary_lines

    rescued, remaining, done = [], [], False
    for line in summary_lines:
        stripped = line.strip()
        if done:
            remaining.append(line); continue
        if not stripped:
            rescued.append(line); continue
        words = stripped.split()
        is_name = (
            2 <= len(words) <= 5
            and re.match(r"^[A-Za-z]", stripped)
            and re.match(r"^[A-Za-z\s.\-']+$", stripped)
            and not re.search(r"\d", stripped)
            and len(stripped) <= 70
        )
        is_contact = bool(
            EMAIL_RE.search(stripped) or PHONE_RE.search(stripped) or LINKEDIN_RE.search(stripped)
        )
        if is_name or is_contact:
            rescued.append(line)
        else:
            done = True
            remaining.append(line)

    if rescued:
        salvaged = extract_contact_info(rescued)
        if salvaged.get("name") and not contact.get("name"):
            contact["name"]  = salvaged["name"]
        if salvaged.get("email") and not contact.get("email"):
            contact["email"] = salvaged["email"]
        if salvaged.get("phone") and not contact.get("phone"):
            contact["phone"] = salvaged["phone"]
    return contact, remaining


def _strip_name_from_summary(summary: str, contact: Dict) -> str:
    """Remove candidate name / email that accidentally starts the summary string."""
    name = contact.get("name", "").strip()
    if name:
        summary = re.sub(
            r"^" + re.escape(name) + r"[\s|,.-]*", "", summary, flags=re.IGNORECASE
        ).strip()
    email = contact.get("email", "").strip()
    if email:
        summary = re.sub(
            r"^" + re.escape(email) + r"[\s|,.-]*", "", summary, flags=re.IGNORECASE
        ).strip()
    return summary


# ─────────────────────────────────────────────────────────────────────────────

def parse_skill_set(file_path: str, file_type: str) -> Dict[str, List[str]]:
    """Parse and return only the skills section from a resume file."""
    raw_text = extract_text(file_path, file_type)
    sections = split_sections(raw_text)
    return parse_skills(sections.get("skills", []))


# ─────────────────────────────────────────────────────────────────────────────

def parse_resume(file_path: str, file_type: str) -> Dict:
    raw_text = extract_text(file_path, file_type)
    sections = split_sections(raw_text)

    contact = extract_contact_info(sections.get("header", []))
    summary_lines = sections.get("summary", [])

    if not summary_lines:
        summary_lines = _extract_summary_from_raw_text(raw_text)
    if not summary_lines:
        summary_lines = _extract_summary_from_header_lines(sections.get("header", []))

    # Rescue name/contact if they landed in the summary block
    contact, summary_lines = _salvage_contact_from_summary(contact, summary_lines)

    bullet_line_re = re.compile(r"^\s*(?:[•\u2022\u2023\u25E6\u2043\u2219\-\*\u2192\u25BA\uf0b7]+|\d+[.)])\s*")
    inline_bullet_re = re.compile(r"\s*[•\u2022\u2023\u25E6\u2043\u2219\-\*\u2192\u25BA\uf0b7]+\s+")

    summary_points: List[str] = []
    summary_plain: List[str] = []
    for line in summary_lines:
        text = line.strip()
        if not text:
            continue

        if bullet_line_re.match(text):
            cleaned = bullet_line_re.sub("", text).strip()
            if cleaned:
                summary_points.append(cleaned)
            continue

        if inline_bullet_re.search(text):
            parts = [p.strip() for p in inline_bullet_re.split(text) if p.strip()]
            summary_points.extend(parts)
            continue

        summary_plain.append(text)

    if summary_points:
        summary = "\n".join(summary_points)
    else:
        summary = " ".join(summary_plain)

    summary = _strip_name_from_summary(summary, contact)

    skills = parse_skills(sections.get("skills", []))
    experience = parse_experience(sections.get("experience", []))
    education = parse_education(sections.get("education", []))
    certifications = parse_list_section(sections.get("certifications", []))
    projects = parse_projects(sections.get("projects", []))
    achievements = parse_list_section(sections.get("achievements", []))

    return {
        "contact": contact,
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "education": education,
        "certifications": certifications,
        "projects": projects,
        "achievements": achievements,
    }
