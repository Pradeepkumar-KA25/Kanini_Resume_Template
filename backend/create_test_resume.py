from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('Priya Ramesh', 0)
doc.add_paragraph('priya.ramesh@email.com | +91-9876543210 | Chennai, India')
doc.add_paragraph('linkedin.com/in/priyaramesh | github.com/priyaramesh')

doc.add_heading('Professional Summary', 1)
doc.add_paragraph('Results-driven Java Backend Developer with 5+ years of experience building scalable microservices and REST APIs. Expertise in Spring Boot, AWS, and CI/CD pipelines.')

doc.add_heading('Technical Skills', 1)
doc.add_paragraph('Programming Languages: Java, Python, TypeScript')
doc.add_paragraph('Web & Frameworks: Spring Boot, Angular, Node.js, REST')
doc.add_paragraph('Cloud & DevOps: AWS, Docker, Jenkins, CI/CD')
doc.add_paragraph('Databases: PostgreSQL, MySQL, MongoDB, Redis')
doc.add_paragraph('Tools: Git, GitHub, JIRA, Postman, Swagger')

doc.add_heading('Professional Experience', 1)
doc.add_paragraph('Kanini Technologies | Senior Software Engineer | Chennai, India')
doc.add_paragraph('Jan 2022 - Present')
doc.add_paragraph('• Designed and implemented 15+ REST APIs using Spring Boot')
doc.add_paragraph('• Reduced API response time by 40% through Redis caching')
doc.add_paragraph('• Led a team of 4 developers, conducted code reviews')
doc.add_paragraph('• Deployed services on AWS EC2 and RDS using CI/CD pipelines')
doc.add_paragraph('')
doc.add_paragraph('Infosys | Software Engineer | Bangalore, India')
doc.add_paragraph('Jun 2019 - Dec 2021')
doc.add_paragraph('• Built microservices for banking application using Java and Spring Boot')
doc.add_paragraph('• Integrated PostgreSQL with JPA/Hibernate for data persistence')
doc.add_paragraph('• Wrote unit and integration tests using JUnit and Mockito')

doc.add_heading('Key Projects', 1)
doc.add_paragraph('E-Commerce Order Management System')
doc.add_paragraph('• Built event-driven order processing using Apache Kafka and Spring Boot')
doc.add_paragraph('• Technologies: Java, Spring Boot, Kafka, PostgreSQL, Docker, AWS')
doc.add_paragraph('')
doc.add_paragraph('Healthcare Patient Portal')
doc.add_paragraph('• Developed RESTful APIs for patient record management')
doc.add_paragraph('• Technologies: Node.js, MongoDB, Angular, Redis')

doc.add_heading('Education', 1)
doc.add_paragraph('B.E. Computer Science | Anna University | 2019')

doc.add_heading('Certifications', 1)
doc.add_paragraph('AWS Certified Solutions Architect - Associate (2023)')
doc.add_paragraph('Oracle Certified Professional Java SE 11')

doc.save('realistic_resume.docx')
print('Created realistic_resume.docx')
