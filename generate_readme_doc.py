"""Convert frontend-ng/README.md to README.docx (plain Word document)."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x1B, 0x3A, 0x6B)
ORANGE = RGBColor(0xE8, 0x77, 0x22)
GRAY   = RGBColor(0x55, 0x65, 0x77)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
CODE_BG = "F4F4F4"

README_PATH = Path(__file__).parent / "frontend-ng" / "README.md"
OUTPUT_PATH = Path(__file__).parent / "Kanini_Resume_Builder_README.docx"


# ── helpers ──────────────────────────────────────────────────────────────────

def set_para_spacing(para, before=0, after=4):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"),  str(after  * 20))
    pPr.append(spacing)


def add_run(para, text, bold=False, italic=False, sz=11, color=None, font="Calibri"):
    r = para.add_run(text)
    r.font.name  = font
    r.font.size  = Pt(sz)
    r.font.bold  = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def add_shaded_block(doc, lines):
    """Render a code block as a shaded, monospace paragraph block."""
    for line in lines:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        # shade background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  CODE_BG)
        pPr.append(shd)
        # indent
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"),  "360")
        ind.set(qn("w:right"), "360")
        pPr.append(ind)
    # small gap after block
    gap = doc.add_paragraph()
    set_para_spacing(gap, before=0, after=4)


def render_inline(para, text, base_sz=11):
    """Render a line that may contain **bold** and `code` spans."""
    # split on **bold** and `code`
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            add_run(para, tok[2:-2], bold=True, sz=base_sz, color=BLACK)
        elif tok.startswith("`") and tok.endswith("`"):
            r = para.add_run(tok[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(base_sz - 1)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            # strip markdown links [text](url) → just "text"
            tok = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', tok)
            if tok:
                add_run(para, tok, sz=base_sz, color=BLACK)


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()
sec = doc.sections[0]
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

# default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# ── parse README ──────────────────────────────────────────────────────────────

lines = README_PATH.read_text(encoding="utf-8").splitlines()

in_code   = False
code_lines = []
i = 0

while i < len(lines):
    line = lines[i]

    # ── code fence ────────────────────────────────────────────────────────────
    if line.strip().startswith("```"):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            add_shaded_block(doc, code_lines)
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # ── horizontal rule ───────────────────────────────────────────────────────
    if re.match(r'^---+$', line.strip()):
        i += 1
        continue

    # ── H1 ────────────────────────────────────────────────────────────────────
    if line.startswith("# ") and not line.startswith("## "):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=6)
        text = line[2:].strip()
        r = p.add_run(text)
        r.font.name  = "Calibri"
        r.font.size  = Pt(20)
        r.font.bold  = True
        r.font.color.rgb = NAVY
        i += 1
        continue

    # ── H2 ────────────────────────────────────────────────────────────────────
    if line.startswith("## "):
        p = doc.add_paragraph()
        set_para_spacing(p, before=10, after=4)
        r = p.add_run(line[3:].strip())
        r.font.name  = "Calibri"
        r.font.size  = Pt(14)
        r.font.bold  = True
        r.font.color.rgb = NAVY
        i += 1
        continue

    # ── H3 ────────────────────────────────────────────────────────────────────
    if line.startswith("### "):
        p = doc.add_paragraph()
        set_para_spacing(p, before=6, after=2)
        r = p.add_run(line[4:].strip())
        r.font.name  = "Calibri"
        r.font.size  = Pt(12)
        r.font.bold  = True
        r.font.color.rgb = ORANGE
        i += 1
        continue

    # ── bullet / sub-bullet ───────────────────────────────────────────────────
    if re.match(r'^(\s{2,4})?[-*] ', line):
        indent = len(line) - len(line.lstrip())
        bullet_text = re.sub(r'^(\s*[-*] )', '', line)
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=2)
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        left = "720" if indent >= 2 else "360"
        ind.set(qn("w:left"),    left)
        ind.set(qn("w:hanging"), "180")
        pPr.append(ind)
        bullet_run = p.add_run("•  ")
        bullet_run.font.color.rgb = ORANGE
        bullet_run.font.bold = True
        bullet_run.font.size = Pt(11)
        render_inline(p, bullet_text, base_sz=11)
        i += 1
        continue

    # ── blank line ────────────────────────────────────────────────────────────
    if line.strip() == "":
        i += 1
        continue

    # ── normal paragraph ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=4)
    render_inline(p, line.strip(), base_sz=11)
    i += 1


# ── save ─────────────────────────────────────────────────────────────────────

doc.save(str(OUTPUT_PATH))
print(f"Saved: {OUTPUT_PATH}")
