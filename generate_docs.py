"""Generate DOCUMENTATION.docx for Kanini Resume Builder POC."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x1B, 0x3A, 0x6B)
ORANGE = RGBColor(0xE8, 0x77, 0x22)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x71, 0x80, 0x96)
LIGHT  = RGBColor(0xE8, 0xF0, 0xFB)
BLACK  = RGBColor(0x00, 0x00, 0x00)

doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin  = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)


# ── helpers ──────────────────────────────────────────────────────────────────

def add_run(para, text, bold=False, italic=False, sz=10, color=None, font="Calibri"):
    r = para.add_run(text)
    r.font.name = font
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r

def set_cell_bg(cell, rgb_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex)
    tcPr.append(shd)

def remove_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for b in ["top","left","bottom","right","insideH","insideV"]:
        el = OxmlElement(f"w:{b}")
        el.set(qn("w:val"),   "none"); el.set(qn("w:sz"), "0")
        el.set(qn("w:space"),"0");     el.set(qn("w:color"),"auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_bottom_border(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single"); bot.set(qn("w:sz"),  "8")
    bot.set(qn("w:space"), "2");      bot.set(qn("w:color"), hex_color)
    pBdr.append(bot); pPr.append(pBdr)

def h1(text):
    """Section heading: navy bold 14pt with orange bottom rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, bold=True, sz=14, color=NAVY)
    add_bottom_border(p, "E87722")
    return p

def h2(text):
    """Sub-heading: navy bold 11pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, bold=True, sz=11, color=NAVY)
    return p

def body(text, sz=10, italic=False, color=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    add_run(p, text, sz=sz, italic=italic, color=color or BLACK)
    return p

def bullet(text, sz=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, sz=sz)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    add_run(p, text, sz=9, font="Courier New", color=GRAY)
    return p

def table_2col(rows_data, hdr=None):
    """Simple 2-column table. rows_data = list of (col1, col2)."""
    n = len(rows_data) + (1 if hdr else 0)
    tbl = doc.add_table(rows=n, cols=2)
    remove_borders(tbl)
    r_idx = 0
    if hdr:
        for ci, h_txt in enumerate(hdr):
            cell = tbl.rows[0].cells[ci]
            set_cell_bg(cell, "1B3A6B")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Inches(0.05)
            add_run(p, h_txt, bold=True, sz=9, color=WHITE)
        r_idx = 1
    for col1, col2 in rows_data:
        row = tbl.rows[r_idx]
        for ci, (val, is_bold) in enumerate([(col1, True), (col2, False)]):
            cell = row.cells[ci]
            set_cell_bg(cell, "F0F4FA" if r_idx % 2 == (1 if hdr else 0) else "FFFFFF")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Inches(0.05)
            add_run(p, val, bold=is_bold, sz=9)
        r_idx += 1
    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(3.8)
    return tbl


# ═══════════════════════════════════════════════════════════════════════════════
# COVER — full-width navy header
# ═══════════════════════════════════════════════════════════════════════════════
cover_tbl = doc.add_table(rows=1, cols=1)
remove_borders(cover_tbl)
cover_cell = cover_tbl.rows[0].cells[0]
set_cell_bg(cover_cell, "1B3A6B")

cp1 = cover_cell.paragraphs[0]
cp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp1.paragraph_format.space_before = Pt(30)
add_run(cp1, "KANINI", bold=True, sz=28, color=WHITE)
add_run(cp1, "  Resume Builder", bold=False, sz=22, color=RGBColor(0xA0, 0xB4, 0xCC))

cp2 = cover_cell.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp2.paragraph_format.space_before = Pt(4)
add_run(cp2, "Proof of Concept — Technical Documentation", sz=11, color=RGBColor(0xBF, 0xD4, 0xEE))

cp3 = cover_cell.add_paragraph()
cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp3.paragraph_format.space_before = Pt(10)
cp3.paragraph_format.space_after  = Pt(30)
add_run(cp3, "Kanini Software Solutions  ·  July 2026", sz=9,
        italic=True, color=RGBColor(0x71, 0x80, 0x96))

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Overview")
body("The Kanini Resume Builder is a web application that automatically converts any uploaded resume "
     "(PDF or Word document) into two standardised, professionally formatted Word documents matching "
     "Kanini's official resume templates.")
doc.add_paragraph()

table_2col([
    ("Template 1 — Kanini Format ",  "Internal Kanini standard format (Sample KANINI Profile)"),
    ("Template 2 — Deloitte Format Profile", "Client submission format (Kanini Format )"),
], hdr=["Template", "Purpose"])

doc.add_paragraph()
body("The application uses OpenAI GPT-4o-mini to intelligently extract all resume information "
     "and populate both templates accurately and consistently.")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. END-TO-END FLOW
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. How It Works — End-to-End Flow")

steps = [
    ("Step 1", "User uploads a PDF or DOCX resume via the web interface (drag & drop or browse)."),
    ("Step 2", "Backend extracts raw text from the uploaded file."),
    ("Step 3", "OpenAI GPT-4o-mini parses the text and returns structured JSON (contact, summary, skills, experience, education, projects)."),
    ("Step 4", "Two Word documents are generated — kanini_classic.docx and deloitte_format.docx."),
    ("Step 5", "Scaled HTML previews of both templates are rendered side-by-side in the browser."),
    ("Step 6", "User downloads either template as .docx (Word) or .pdf."),
]
table_2col(steps, hdr=["Step", "Description"])


# ═══════════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Technology Stack")
table_2col([
    ("Frontend Framework",   "Angular 19"),
    ("Backend Framework",    "Python FastAPI 0.104"),
    ("AI / NLP",             "OpenAI GPT-4o-mini (API v1)"),
    ("Word Generation",      "python-docx 1.1.0"),
    ("PDF Text Extraction",  "pdfminer.six 20221105"),
    ("DOCX Text Extraction", "docx2txt 0.8"),
    ("Application Server",   "Uvicorn 0.24 (ASGI)"),
    ("API Communication",    "REST / JSON over HTTP"),
], hdr=["Layer", "Technology"])


# ═══════════════════════════════════════════════════════════════════════════════
# 4. PROJECT STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. Project Structure")
for line in [
    "Kanini_Template_POC/",
    "  ├── backend/",
    "  │   ├── main.py               ← API routes (upload, download, health)",
    "  │   ├── ai_parser.py          ← OpenAI GPT-4o-mini integration",
    "  │   ├── resume_parser.py      ← Regex-based fallback parser",
    "  │   ├── template_generator.py ← Word doc & HTML preview generators",
    "  │   ├── requirements.txt      ← Python dependencies",
    "  │   └── .env                  ← OPENAI_API_KEY (not in source control)",
    "  │",
    "  ├── frontend-ng/              ← Angular 19 frontend",
    "  │   └── src/app/",
    "  │       ├── components/",
    "  │       │   ├── file-upload/  ← Drag & drop upload UI",
    "  │       │   ├── loading-view/ ← Processing spinner",
    "  │       │   ├── results-view/ ← Side-by-side template previews",
    "  │       │   └── template-card/← Preview card + download buttons",
    "  │       ├── services/         ← HTTP calls to backend",
    "  │       └── models/           ← TypeScript data interfaces",
    "  │",
    "  ├── setup.bat                 ← One-click environment setup",
    "  └── start.bat                 ← Start both servers",
]:
    code_block(line)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. API ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. API Endpoints")
table_2col([
    ("GET  /api/health",                                  "Backend health check — returns {\"status\":\"ok\"}"),
    ("POST /api/upload",                                  "Upload resume file → returns parsed data + HTML previews"),
    ("GET  /api/download/{session}/{template}/{format}",  "Download generated file (docx or pdf)"),
], hdr=["Endpoint", "Description"])

doc.add_paragraph()
h2("Upload Response Schema")
for line in [
    '{',
    '  "session_id":  "uuid-string",',
    '  "resume_data": {',
    '    "contact":        { name, email, phone, location, linkedin },',
    '    "summary":        "Professional summary...",',
    '    "skills":         { "Category": ["skill1", "skill2"] },',
    '    "experience":     [{ title, company, dates, responsibilities:[] }],',
    '    "education":      [{ degree, institution, year }],',
    '    "certifications": [],',
    '    "projects":       [{ name, description, technologies:[] }],',
    '    "achievements":   []',
    '  },',
    '  "preview_html": { "template1": "<html>", "template2": "<html>" }',
    '}',
]:
    code_block(line)


# ═══════════════════════════════════════════════════════════════════════════════
# 6. AI PARSING
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. AI Parsing — OpenAI Integration")
body("File: backend/ai_parser.py", italic=True, color=GRAY)
doc.add_paragraph()
table_2col([
    ("Model Used",         "GPT-4o-mini"),
    ("Response Format",    "JSON object (structured, reliable output)"),
    ("Max Input Length",   "14,000 characters per request"),
    ("Approx. Token Cost", "1,000 – 3,000 tokens per resume upload"),
    ("Fallback",           "If API key missing or call fails → regex parser"),
], hdr=["Property", "Value"])

doc.add_paragraph()
h2("Fallback Chain")
bullet("API key present → GPT-4o-mini parses resume → structured JSON")
bullet("AI call fails → regex-based resume_parser.py is used instead")
bullet("Regex also fails → HTTP 422 error returned to user")

doc.add_paragraph()
h2("API Key Configuration")
body("Add the OpenAI API key to backend/.env:")
code_block("OPENAI_API_KEY=sk-...your-key-here...")


# ═══════════════════════════════════════════════════════════════════════════════
# 7. TEMPLATE FORMATS
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Template Formats")

h2("Template 1 — Kanini Format ")
body("Matches the exact structure of Sample KANINI Profile.docx:")
for line in [
    "CANDIDATE NAME  (bold, 14pt, Calibri)",
    "Mobile No: xxx  |  Email Id: xxx",
    "",
    "Profile Summary  (bold header)",
    "   • Summary bullet 1",
    "   • Summary bullet 2",
    "",
    "Technical Skills:  (bold header)",
    "   • Category: skill1, skill2   ← entire line bold",
    "",
    "Work Experience:  (bold header)",
    "   Company Name  :  Company XYZ",
    "   Designation   :  Job Title",
    "   Duration      :  Jan 2020 – Present",
    "",
    "Project Summary:                ← combined with Project I: in one paragraph",
    "",
    "Project I:",
    "   Job Title | Company XYZ",
    "   Roles and Responsibilities:",
    "   • Responsibility 1",
    "",
    "Educational Qualification:",
    "   Degree (Year) from Institution",
]:
    code_block(line)

doc.add_paragraph()
h2("Template 2 — Deloitte Format (Kanini Profile Format)")
body("Matches the exact structure of Kanini Format .docx:")
for line in [
    "                              CANDIDATE NAME  ← right-aligned, 10pt bold",
    "",
    "Professional Summary:   ← Heading 1 style",
    "Summary sentence 1.",
    "",
    "Technical Skills:   ← Heading 1 style",
    "Programming         : React JS, JavaScript, C#, .NET",
    "Databases           : MS SQL Server",
    "",
    "Working Experience:   ← Heading 1 style",
    "Company Name        : KANINI SOFTWARE SOLUTIONS",
    "Designation         : Junior Associate",
    "Duration            : July 2022 - Till date",
    "",
    "Project Summary:   ← Heading 1 style",
    "Project – I  (bold)",
    "Client              : Deloitte - ESG",
    "Technologies        : React, Redux, JavaScript, CSS, Azure",
    "Role                : Developer",
    "",
    "Description of Project:   ← Heading 1 style",
    "Project description text...",
    "",
    "Roles and Responsibilities:   ← Heading 1 style",
    "Responsibility 1",
    "Responsibility 2",
    "",
    "EDUCATIONAL QUALIFICATION:   ← Heading 1 style",
    "BE (2022) from College Name",
]:
    code_block(line)


# ═══════════════════════════════════════════════════════════════════════════════
# 8. SUPPORTED INPUT FORMATS
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Supported Input Formats")
table_2col([
    ("PDF",       ".pdf — pdfminer.six (text-based PDFs only)"),
    ("Word 2007+",".docx — docx2txt"),
], hdr=["Format", "Details"])
doc.add_paragraph()
bullet("Maximum file size: 10 MB")
bullet("Scanned / image-based PDFs are NOT supported (no OCR)")
bullet("Text must be selectable in the source document")


# ═══════════════════════════════════════════════════════════════════════════════
# 9. FRONTEND FEATURES
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. Frontend Features")
table_2col([
    ("File Upload",         "Drag & drop or click-to-browse; PDF and DOCX accepted"),
    ("Processing Indicator","Animated spinner with live status steps"),
    ("Side-by-Side Preview","Both templates displayed together, scaled to fit screen"),
    ("Candidate Name Bar",  "Candidate name shown inside each template card header"),
    ("Download Buttons",    "One-click download as .docx (Word) or .pdf per template"),
    ("Responsive Layout",   "Stacks to single column on screens narrower than 900px"),
    ("Auto Height Scaling", "Template preview height auto-calculates to show full content"),
], hdr=["Feature", "Description"])


# ═══════════════════════════════════════════════════════════════════════════════
# 10. SETUP & RUNNING
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. Setup & Running")

h2("Prerequisites")
bullet("Python 3.10 or later")
bullet("Node.js 18 or later")
bullet("OpenAI API key (stored in backend/.env)")

doc.add_paragraph()
h2("Quick Start (Windows)")
code_block("# First-time environment setup")
code_block("setup.bat")
doc.add_paragraph()
code_block("# Start both servers (backend port 8000, frontend port 4200)")
code_block("start.bat")
doc.add_paragraph()
code_block("# Open browser at:")
code_block("http://localhost:4200")

doc.add_paragraph()
h2("Manual Start")
body("Backend:")
code_block("cd backend")
code_block("python -m uvicorn main:app --port 8000")
doc.add_paragraph()
body("Frontend:")
code_block("cd frontend-ng")
code_block("ng serve --port 4200")


# ═══════════════════════════════════════════════════════════════════════════════
# 11. KEY DEPENDENCIES
# ═══════════════════════════════════════════════════════════════════════════════
h1("11. Key Dependencies")
h2("Backend (Python)")
table_2col([
    ("fastapi==0.104.1",        "Web API framework"),
    ("uvicorn==0.24.0",         "ASGI application server"),
    ("python-docx==1.1.0",      "Word document generation"),
    ("pdfminer.six==20221105",  "PDF text extraction"),
    ("openai>=1.0.0",           "OpenAI GPT-4o-mini API client"),
    ("python-dotenv>=1.0.0",    "Load API key from .env file"),
    ("docx2txt==0.8",           "DOCX text extraction"),
], hdr=["Package", "Purpose"])

doc.add_paragraph()
h2("Frontend (Node / Angular)")
table_2col([
    ("@angular/core ^19.0.0",   "UI component framework"),
    ("@angular/common/http",    "HTTP client for backend API calls"),
    ("TypeScript 5.x",          "Type-safe JavaScript"),
], hdr=["Package", "Purpose"])


# ═══════════════════════════════════════════════════════════════════════════════
# 12. SECURITY NOTES
# ═══════════════════════════════════════════════════════════════════════════════
h1("12. Security Notes")
table_2col([
    ("API Key Storage",  "Stored in backend/.env — not committed to source control"),
    ("Session Isolation","Each upload gets a UUID session; files are stored in OS temp directory"),
    ("File Validation",  "Only .pdf and .docx extensions accepted; max 10 MB"),
    ("CORS Policy",      "All origins allowed (POC setting; should be restricted in production)"),
    ("No Authentication","POC only — no user login or access control"),
], hdr=["Area", "Note"])


# ═══════════════════════════════════════════════════════════════════════════════
# 13. LIMITATIONS (POC SCOPE)
# ═══════════════════════════════════════════════════════════════════════════════
h1("13. Limitations (POC Scope)")
table_2col([
    ("No User Authentication",    "Sessions are UUID-only; no login required"),
    ("Ephemeral File Storage",    "Generated files are lost when the server restarts"),
    ("OpenAI API Cost",           "Approximately 1,000–3,000 GPT tokens per upload"),
    ("Scanned PDFs Not Supported","OCR is not included; requires text-based PDFs"),
    ("No History / Dashboard",    "No record of previous uploads"),
    ("Single-Machine Deployment", "Not yet containerised or cloud-deployed"),
], hdr=["Limitation", "Detail"])

doc.add_paragraph()
body("These limitations are appropriate for a POC and would be addressed before production deployment.",
     italic=True, color=GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
out = r"d:\Kanini_Template_POC\Kanini_Resume_Builder_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
